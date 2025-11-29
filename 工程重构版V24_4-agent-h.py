import argparse
import subprocess
import time
import requests
import re
import json
import os
import sys
import random
import shutil
import glob
import traceback
import hashlib
import base64
from io import BytesIO
import matplotlib
matplotlib.use('Agg')  # 非交互式后端，避免显示窗口
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from datetime import datetime

# 优先尝试使用 LangChain 的标题/递归切分，如缺失则降级
try:
    from langchain_text_splitters import MarkdownHeaderTextSplitter, RecursiveCharacterTextSplitter
    from langchain.schema import Document
    HAS_LANGCHAIN = True
except ImportError:
    try:
        from langchain.text_splitter import MarkdownHeaderTextSplitter, RecursiveCharacterTextSplitter
        from langchain.schema import Document
        HAS_LANGCHAIN = True
    except ImportError:
        HAS_LANGCHAIN = False
        print("⚠️ 未检测到 langchain/text-splitters，分块将回退为简易字符切分。建议: pip install langchain langchain-text-splitters")

# ================== 🎨 字体管理 (解决中文乱码) ==================

def get_chinese_font():
    """动态查找系统中可用的中文字体"""
    # 常见的中文字体路径 (macOS, Linux, Windows)
    font_candidates = [
        # macOS
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/STHeiti Medium.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        # Linux
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
        # Windows
        "C:\\Windows\\Fonts\\simhei.ttf",
        "C:\\Windows\\Fonts\\msyh.ttc",
    ]
    
    # 1. 检查文件是否存在
    for path in font_candidates:
        if os.path.exists(path):
            return fm.FontProperties(fname=path)
            
    # 2. 如果没找到文件，尝试通过 family name 获取 (Matplotlib 默认机制)
    common_families = ['SimHei', 'Arial Unicode MS', 'PingFang SC', 'Heiti TC', 'Microsoft YaHei', 'WenQuanYi Micro Hei']
    for family in common_families:
        try:
            # 尝试加载字体
            fm.findfont(family, fallback_to_default=False)
            return fm.FontProperties(family=family)
        except:
            continue
            
    return None

# 初始化字体
CHINESE_FONT = get_chinese_font()
if CHINESE_FONT:
    # 全局设置
    try:
        font_name = CHINESE_FONT.get_name() if hasattr(CHINESE_FONT, 'get_name') else 'Unknown'
        plt.rcParams['font.family'] = font_name
    except:
        # 如果是路径创建的 FontProperties，get_name 可能返回文件名，这里做容错
        font_name = 'Loaded'
    print(f"✅ 已加载中文字体: {font_name}")
else:
    print("⚠️ 未找到合适的中文字体，图表中文可能显示乱码")

plt.rcParams['axes.unicode_minus'] = False  # 处理负号显示

# ================== ⚙️ 配置管理类 (解决硬编码问题) ==================

class Config:
    def __init__(self):
        # 优先从环境变量读取，否则设为 None
        self.GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
        self.TAVILY_API_KEY = os.getenv("TAVILY_API_KEY", "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")
        
        # 默认配置
        self.GEMINI_OUTLINE_MODEL = "gemini-3-pro-preview"  # 升级：使用最新 Gemini 3.0 Pro 构建大纲
        self.GEMINI_GEN_MODEL = "gemini-3-pro-preview"     # 升级：使用最新 Gemini 3.0 Pro 确保最佳性能
        self.GEMINI_PRO_MODEL = os.getenv("GEMINI_PRO_MODEL", "gemini-3-pro-preview")   # Pro 角色（规划/润色/代码）
        self.GEMINI_FLASH_MODEL = os.getenv("GEMINI_FLASH_MODEL", "gemini-2.5-flash")   # Flash 角色（搜索/粗写/数据整理）
        self.GEMINI_FLASH_LITE_MODEL = os.getenv("GEMINI_FLASH_LITE_MODEL", "gemini-2.5-flash-lite")  # 极致省钱版
        
        # 路径配置 (使用相对路径或用户目录，避免硬编码绝对路径)
        self.BASE_DIR = os.path.join(os.path.expanduser("~"), "Research_Workspace")
        
        # 大纲结构配置 (可动态调整)
        self.OUTLINE_CHAPTERS = int(os.getenv("OUTLINE_CHAPTERS", 8))  # 章节数
        self.OUTLINE_SECTIONS = int(os.getenv("OUTLINE_SECTIONS", 3))  # 每章小节数
        self.OUTLINE_SUBSECTIONS = int(os.getenv("OUTLINE_SUBSECTIONS", 3))  # 每小节子节数
        
        # 目标页数配置 (自动计算章节结构)
        self.TARGET_PAGES = int(os.getenv("TARGET_PAGES", 0))  # 0 表示不自动调整
        self.WORDS_PER_PAGE = int(os.getenv("WORDS_PER_PAGE", 500))  # 每页平均字数（含格式）
        
        # 写作权重配置
        self.SECTION_WEIGHT = float(os.getenv("SECTION_WEIGHT", 0.5))  # 小节权重（vs 子节）
        self.QUALITY_THRESHOLD = float(os.getenv("QUALITY_THRESHOLD", 8.0))  # 提高为8.0，确保高质量
        self.MAX_REFINEMENT_ROUNDS = int(os.getenv("MAX_REFINEMENT_ROUNDS", 4))  # 增加到4轮改进
        self.USE_TFIDF_RAG = os.getenv("USE_TFIDF_RAG", "true").lower() == "true"  # 启用向量检索

        # 代理配置
        self.PROXY_URL = os.getenv("HTTP_PROXY") or "http://127.0.0.1:6152"
        self.PROXIES_CLOUD = {"http": self.PROXY_URL, "https": self.PROXY_URL}
        self.PROXIES_LOCAL = {"http": None, "https": None} # 强制直连

        # MinerU 配置（用于更快的素材解析）
        self.USE_MINERU = os.getenv("USE_MINERU", "true").lower() == "true"
        default_mineru_home = os.path.expanduser("~/mineru")
        default_mineru_cmds = [
            os.path.join(default_mineru_home, "bin", "mineru"),                   # 本机安装 (macOS M2)
            "/opt/anaconda3/envs/mineru_env/bin/mineru",                         # 备选: anaconda 环境
            "/opt/anaconda3/envs/mineru262-env/bin/mineru"                       # 旧路径兼容
        ]
        env_mineru_cmd = os.getenv("MINERU_CMD")
        candidate_cmds = []
        if env_mineru_cmd:
            candidate_cmds.append(os.path.expanduser(env_mineru_cmd))
        candidate_cmds.extend(default_mineru_cmds)
        resolved_cmd = None
        for candidate in candidate_cmds:
            expanded = os.path.expanduser(candidate)
            if os.path.exists(expanded):
                resolved_cmd = expanded
                break
        if not resolved_cmd:
            print("⚠️ 未在预设路径找到 MinerU，将尝试使用系统路径中的 'mineru'")
            resolved_cmd = "mineru"
        self.MINERU_CMD = resolved_cmd
        self.MINERU_IN_DIR = os.path.expanduser(os.getenv("MINERU_IN", os.path.join(default_mineru_home, "workflow", "in")))
        self.MINERU_OUT_DIR = os.path.expanduser(os.getenv("MINERU_OUT", os.path.join(default_mineru_home, "workflow", "out")))
        self.MINERU_METHOD = os.getenv("MINERU_METHOD", "auto")
        self.MINERU_BACKEND = os.getenv("MINERU_BACKEND", "pipeline")
        self.MINERU_LANG = os.getenv("MINERU_LANG", "ch")
        self.MINERU_TIMEOUT = int(os.getenv("MINERU_TIMEOUT", 600))

    def validate(self):
        """启动前自检"""
        errors = []
        if not self.GEMINI_API_KEY or "Your_Key" in self.GEMINI_API_KEY:
            # 交互式补救
            self.GEMINI_API_KEY = input("⚠️ 未检测到 GEMINI_API_KEY，请输入: ").strip()
            if not self.GEMINI_API_KEY: errors.append("缺少 GEMINI_API_KEY")
            
        if not self.TAVILY_API_KEY or "Your_Key" in self.TAVILY_API_KEY:
            print("ℹ️ 未检测到 TAVILY_API_KEY，联网搜索功能将受限。")
            
        if errors:
            print("\n❌ 配置错误，无法启动:")
            for e in errors: print(f"  - {e}")
            sys.exit(1)
    
    def calculate_outline_structure(self):
        """
        根据目标页数自动计算大纲结构
        返回 (chapters, sections, subsections)
        """
        if self.TARGET_PAGES <= 0:
            # 不自动调整，使用手动配置
            return self.OUTLINE_CHAPTERS, self.OUTLINE_SECTIONS, self.OUTLINE_SUBSECTIONS
        
        # 反推需要的总字数
        total_words_needed = self.TARGET_PAGES * self.WORDS_PER_PAGE
        
        # 每个子节平均字数
        words_per_subsection = 1000
        
        # 需要的子节总数
        total_subsections_needed = total_words_needed // words_per_subsection
        
        # 根据子节数反推章节结构
        # 策略：优先增加章节数，其次增加小节数，最后增加子节数
        
        # 假设基础配置
        chapters = max(5, total_subsections_needed // 12)  # 至少 5 章
        sections = 3
        subsections = max(3, total_subsections_needed // (chapters * sections))
        
        # 微调
        actual_subsections = chapters * sections * subsections
        if actual_subsections < total_subsections_needed:
            subsections += 1
        
        return chapters, sections, subsections
    
    def estimate_page_count(self):
        """估算报告总页数"""
        total_subsections = self.OUTLINE_CHAPTERS * self.OUTLINE_SECTIONS * self.OUTLINE_SUBSECTIONS
        estimated_words = total_subsections * 1000  # 每个子节 1000 字
        estimated_pages = estimated_words / self.WORDS_PER_PAGE
        return estimated_pages

# 初始化配置单例
CONF = Config()

# ================== 📚 RAG 知识库 (修复中文检索) ==================

class MaterialManager:
    def __init__(self, folder):
        self.chunks = []
        self.folder = folder
        self.failed_files = []
        self.chunk_stats = {}  # 记录每个文件的 chunk 统计
        self.use_mineru = getattr(CONF, "USE_MINERU", False)
        self.mineru_cmd = getattr(CONF, "MINERU_CMD", "")
        self.mineru_in_dir = getattr(CONF, "MINERU_IN_DIR", "")
        self.mineru_out_dir = getattr(CONF, "MINERU_OUT_DIR", "")
        self.mineru_method = getattr(CONF, "MINERU_METHOD", "auto")
        self.mineru_backend = getattr(CONF, "MINERU_BACKEND", "pipeline")
        self.mineru_lang = getattr(CONF, "MINERU_LANG", "ch")
        self.mineru_timeout = getattr(CONF, "MINERU_TIMEOUT", 600)
        self.use_tfidf = getattr(CONF, "USE_TFIDF_RAG", False)
        self.vectorizer = None
        self.tfidf_matrix = None
        self.load()

    def load(self):
        if not os.path.exists(self.folder):
            try:
                os.makedirs(self.folder)
                print(f"📂 [初始化] 已创建素材目录: {self.folder}")
                print("👉 请放入 PDF/Word/Txt 资料，脚本将自动扫描...")
                # 这里不阻塞，允许后续流程运行
            except OSError as e:
                print(f"❌ 无法创建目录: {e}")
                return

        # 递归获取所有子目录中的文件
        files = []
        for root, dirs, filenames in os.walk(self.folder):
            for fname in filenames:
                files.append(os.path.join(root, fname))
        
        print(f"\n📚 正在扫描素材目录及子目录，发现 {len(files)} 份文件...")
        
        loaded_count = 0
        for fpath in files:
            if os.path.basename(fpath).startswith('.'): continue # 跳过隐藏文件
            try:
                ext = os.path.splitext(fpath)[1].lower()
                content = ""
                source_name = os.path.basename(fpath)
                mineru_used = False
                mineru_md_path = None
                source_type = "standard_loader"
                
                mineru_supported = {'.pdf', '.png', '.jpg', '.jpeg', '.bmp', '.webp', '.tiff', '.gif'}
                if self.use_mineru and ext in mineru_supported:
                    mineru_md_path = self._convert_with_mineru(fpath)
                    if mineru_md_path and os.path.exists(mineru_md_path):
                        try:
                            with open(mineru_md_path, 'r', encoding='utf-8', errors='ignore') as f:
                                content = f.read()
                            mineru_used = True
                            source_name = f"{source_name} [MinerU]"
                            source_type = "mineru_markdown"
                        except Exception as e:
                            print(f"⚠️ MinerU 结果读取失败，改用内置解析: {str(e)[:50]}")
                            content = ""
                
                # 显式依赖检查与错误捕获（MinerU 失败时回退）
                if not content.strip():
                    if ext == '.pdf':
                        try:
                            import pdfplumber
                            with pdfplumber.open(fpath) as pdf:
                                for p in pdf.pages: content += (p.extract_text() or "") + "\n"
                        except ImportError:
                            self.failed_files.append(f"{os.path.basename(fpath)} (缺少 pdfplumber 库)")
                            continue
                            
                    elif ext == '.docx':
                        try:
                            import docx
                            doc = docx.Document(fpath)
                            content = "\n".join([p.text for p in doc.paragraphs])
                        except ImportError:
                            self.failed_files.append(f"{os.path.basename(fpath)} (缺少 python-docx 库)")
                            continue
                        except Exception as e:
                            self.failed_files.append(f"{os.path.basename(fpath)} (读取错误: {str(e)[:50]})")
                            continue
                    
                    elif ext == '.doc':
                        # 旧格式 .doc 文件，尝试用 docx 库兼容模式或提取文本
                        try:
                            import docx
                            doc = docx.Document(fpath)
                            content = "\n".join([p.text for p in doc.paragraphs])
                        except:
                            # 降级处理：直接提取二进制文本
                            try:
                                with open(fpath, 'rb') as f:
                                    raw = f.read()
                                    # 尝试解码为 utf-8 或 gbk
                                    for encoding in ['utf-8', 'gbk', 'latin-1']:
                                        try:
                                            decoded = raw.decode(encoding, errors='ignore')
                                            # 清理控制字符
                                            content = ''.join(c for c in decoded if ord(c) >= 32 or c in '\n\r\t')
                                            break
                                        except:
                                            continue
                            except:
                                content = ""
                        
                        if not content.strip():
                            self.failed_files.append(f"{os.path.basename(fpath)} (无法读取内容)")
                            continue
                            
                    elif ext in ['.txt', '.md']:
                        with open(fpath, 'r', encoding='utf-8', errors='ignore') as f: content = f.read()
                
                if content.strip():
                    # 智能分割：根据内容结构进行更深层次的 chunk (MinerU 优先)
                    file_meta = {"filename": os.path.basename(fpath), "path": fpath}
                    chunks = self.smart_chunk_material(content, source_type, file_meta)
                    
                    for chunk_data in chunks:
                        self.chunks.append(chunk_data)
                    
                    # 记录统计
                    fname = os.path.basename(fpath)
                    self.chunk_stats[fname] = {
                        "chunks": len(chunks),
                        "total_chars": len(content),
                        "status": "成功",
                        "parser": "MinerU" if mineru_used else "builtin"
                    }
                    loaded_count += 1
                else:
                    self.failed_files.append(f"{os.path.basename(fpath)} (内容为空)")
                    
            except Exception as e:
                self.failed_files.append(f"{os.path.basename(fpath)} (读取错误: {str(e)})")

        # 打印加载报告
        total_chunks = len(self.chunks)
        print(f"✅ 成功加载: {loaded_count} 份文件 | 生成 {total_chunks} 个智能 chunk")
        if self.failed_files:
            print("⚠️ 以下文件加载失败:")
            for fail in self.failed_files: print(f"  - {fail}")
        
        # 打印详细统计
        print("\n📊 素材加载统计:")
        for fname, stats in self.chunk_stats.items():
            parser = stats.get("parser", "builtin")
            print(f"  {fname}: {stats['chunks']} chunks ({stats['total_chars']} 字符) [{parser}]")
        
        # 构建 TF-IDF 向量索引（可选）
        self._build_vector_index()

    def _locate_mineru_md(self, base_name):
        """查找 MinerU 已生成的 Markdown，避免重复解析"""
        if not self.mineru_out_dir:
            return None
        candidates = [
            os.path.join(self.mineru_out_dir, f"{base_name}.md"),
            os.path.join(self.mineru_out_dir, base_name, f"{base_name}.md")
        ]
        pattern = os.path.join(self.mineru_out_dir, base_name, "*", f"{base_name}.md")
        candidates.extend(glob.glob(pattern))
        for path in candidates:
            if path and os.path.exists(path):
                return path
        return None

    def _convert_with_mineru(self, fpath):
        """调用 MinerU 将 PDF/图片解析为 Markdown，失败则返回 None"""
        if not self.use_mineru or not self.mineru_cmd or not self.mineru_out_dir:
            return None
        cmd_path = self.mineru_cmd
        if not os.path.exists(cmd_path):
            # 允许走 PATH 中的 mineru 命令
            from shutil import which
            path_cmd = which(cmd_path)
            if not path_cmd:
                return None
            cmd_path = path_cmd
        base_name = os.path.splitext(os.path.basename(fpath))[0]
        cached_md = self._locate_mineru_md(base_name)
        if cached_md:
            return cached_md
        try:
            os.makedirs(self.mineru_out_dir, exist_ok=True)
            if self.mineru_in_dir:
                try:
                    os.makedirs(self.mineru_in_dir, exist_ok=True)
                    dst = os.path.join(self.mineru_in_dir, os.path.basename(fpath))
                    if not os.path.exists(dst):
                        shutil.copy(fpath, dst)
                except Exception:
                    pass
            cmd = [
                cmd_path,
                "-p", fpath,
                "-o", self.mineru_out_dir,
                "-m", self.mineru_method,
                "-b", self.mineru_backend,
                "-l", self.mineru_lang
            ]
            subprocess.run(cmd, check=True, timeout=self.mineru_timeout)
            return self._locate_mineru_md(base_name)
        except Exception as e:
            print(f"⚠️ MinerU 解析失败: {e}")
            return None
    
    def smart_chunk_material(self, content, source_type, file_meta, chunk_size=800, chunk_overlap=100):
        """
        智能分块策略：
        1) MinerU Markdown: 标题切分 + 递归切分，保留 H1/H2/H3 元数据
        2) 其他/失败: 递归字符切分，保留文件元数据
        """
        if not content:
            return []

        def _docs_to_chunks(docs, source_label, base_meta):
            chunk_list = []
            for d in docs:
                text = getattr(d, "page_content", None)
                if text is None and isinstance(d, dict):
                    text = d.get("page_content", "")

                meta = {}
                if hasattr(d, "metadata"):
                    meta = dict(getattr(d, "metadata", {}) or {})
                elif isinstance(d, dict):
                    meta = dict(d.get("metadata", {}))
                merged_meta = {**base_meta, **meta, "source_type": source_label}
                chunk_list.append({
                    "source": base_meta.get("filename", "unknown"),
                    "text": text,
                    "size": len(text),
                    "weight": 1.0,
                    "metadata": merged_meta
                })
            return chunk_list

        # LangChain 不可用时，简易字符切分
        if not HAS_LANGCHAIN:
            chunks = []
            for i in range(0, len(content), chunk_size - chunk_overlap):
                seg = content[i:i+chunk_size]
                chunks.append({
                    "source": file_meta.get("filename", "unknown"),
                    "text": seg,
                    "size": len(seg),
                    "weight": 1.0,
                    "metadata": {**file_meta, "source_type": source_type, "method": "legacy_split"}
                })
            print(f"   🔪 [Legacy Split] {file_meta.get('filename','?')} -> {len(chunks)} chunks")
            return chunks

        # MinerU Markdown: 标题感知切分
        if source_type == "mineru_markdown":
            try:
                headers_to_split_on = [
                    ("#", "Header 1"),
                    ("##", "Header 2"),
                    ("###", "Header 3"),
                ]
                md_splitter = MarkdownHeaderTextSplitter(headers_to_split_on=headers_to_split_on)
                md_docs = md_splitter.split_text(content)

                recursive_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=chunk_size,
                    chunk_overlap=chunk_overlap
                )
                final_docs = recursive_splitter.split_documents(md_docs)
                chunks = _docs_to_chunks(final_docs, "mineru_markdown", file_meta)
                print(f"   🔪 [MinerU Split] {file_meta.get('filename','?')} -> {len(chunks)} chunks")
                return chunks
            except Exception as e:
                print(f"   ⚠️ Markdown 切分失败 ({e})，回退标准切分")
                # 继续走标准切分

        # 标准递归切分
        recursive_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
        std_docs = recursive_splitter.create_documents([content], metadatas=[file_meta])
        chunks = _docs_to_chunks(std_docs, "standard_loader", file_meta)
        print(f"   🔪 [Standard Split] {file_meta.get('filename','?')} -> {len(chunks)} chunks")
        return chunks

    def _build_vector_index(self):
        """可选: 使用 TF-IDF 向量化提升检索效果"""
        if not self.use_tfidf:
            return
        
        # 第一步：尝试导入 sklearn
        try:
            from sklearn.feature_extraction.text import TfidfVectorizer
            has_sklearn = True
        except ImportError as ie:
            print(f"ℹ️ 未安装 sklearn ({ie})，向量检索关闭，继续使用关键词匹配")
            self.use_tfidf = False
            self.vectorizer = None
            self.tfidf_matrix = None
            return
        
        # 第二步：如果成功导入，继续向量化处理
        texts = [c.get("text", "") for c in self.chunks if c.get("text")]
        if not texts:
            print("⚠️ 无有效文本内容用于 TF-IDF 索引")
            self.use_tfidf = False
            self.vectorizer = None
            self.tfidf_matrix = None
            return
        
        try:
            self.vectorizer = TfidfVectorizer(max_features=50000)
            self.tfidf_matrix = self.vectorizer.fit_transform(texts)
            print(f"✅ 已构建 TF-IDF 向量索引: 文档数 {len(texts)}, 维度 {self.tfidf_matrix.shape[1]}")
        except Exception as e:
            print(f"⚠️ TF-IDF 构建失败，回退关键词检索: {type(e).__name__}: {str(e)[:60]}")
            self.vectorizer = None
            self.tfidf_matrix = None
            self.use_tfidf = False

    def retrieve(self, query, top_k=6):
        """支持中文的检索逻辑，带权重计算，并格式化元数据（优先向量检索，回退关键词）"""
        if not self.chunks: 
            return ""
        
        # 修复：使用支持 CJK 的正则进行分词
        # 匹配：连续的汉字 OR 连续的字母数字
        query_words = set(re.findall(r"[\u4e00-\u9fff]+|[a-zA-Z0-9]+", query))
        
        if not query_words and not query:
            return ""

        # 关键词匹配得分
        keyword_scores = {}
        for idx, c in enumerate(self.chunks):
            score = 0
            for w in query_words:
                if w in c['text']:
                    score += 1
            if score > 0:
                keyword_scores[idx] = score * c.get('weight', 1.0)

        combined = []

        # TF-IDF 向量检索（可选，sklearn 可能不可用）
        if self.use_tfidf and self.vectorizer is not None and self.tfidf_matrix is not None:
            try:
                q_vec = self.vectorizer.transform([query])
                if q_vec.nnz > 0:
                    sims = (self.tfidf_matrix @ q_vec.T).toarray().ravel()
                    for idx, sim in enumerate(sims):
                        if sim > 0:
                            kw_bonus = keyword_scores.get(idx, 0) * 0.1  # 关键词轻量加权
                            combined.append((sim + kw_bonus, idx))
                else:
                    # 向量化为空时，回退关键词
                    pass
            except Exception as e:
                # sklearn 相关的任何异常都安全降级
                print(f"ℹ️ 向量检索异常 ({type(e).__name__})，自动回退关键词匹配")
                self.use_tfidf = False
                self.vectorizer = None
                self.tfidf_matrix = None

        # 如果向量检索不可用或未命中，使用关键词匹配
        if not combined:
            combined = [(score, idx) for idx, score in keyword_scores.items()]

        # 排序并截断
        combined.sort(key=lambda x: x[0], reverse=True)

        context = ""
        for rank, (s, idx) in enumerate(combined[:top_k], start=1):
            item = self.chunks[idx]
            meta = item.get("metadata", {}) or {}
            source = meta.get("filename", item.get("source", "unknown"))
            h1 = meta.get("Header 1") or meta.get("H1") or ""
            h2 = meta.get("Header 2") or meta.get("H2") or ""
            h3 = meta.get("Header 3") or meta.get("H3") or ""
            breadcrumb = source
            if h1: breadcrumb += f" > {h1}"
            if h2: breadcrumb += f" > {h2}"
            if h3: breadcrumb += f" > {h3}"
            context += (
                f"\n--- 资料片段 {rank} (匹配度:{s:.3f}) ---\n"
                f"[来源路径]: {breadcrumb}\n"
                f"[内容]: {item['text']}\n"
            )
        
        if context:
            print(f"      🧩 [RAG] 命中 {len(combined[:top_k])} 个片段 (关键词: {list(query_words)[:3]}...)")
        return context

# ================== 🌍 联网搜索 (带容错) ==================

_TAVILY_KEY_WARNED = False

def search_web(query, force=False, cache_dir=None, max_retries=3):
    """
    增强版联网搜索：
    - 稳定缓存 (md5) 跨进程命中
    - 支持代理
    - 指数退避重试
    - 缺失 Key 时仅警告一次
    """
    global _TAVILY_KEY_WARNED

    api_key = os.getenv("TAVILY_API_KEY") or getattr(CONF, "TAVILY_API_KEY", "")
    if not api_key:
        if not _TAVILY_KEY_WARNED:
            print("⚠️ [Search] 未检测到 TAVILY_API_KEY，网络搜索将跳过")
            _TAVILY_KEY_WARNED = True
        return ""
    if not query:
        return ""

    # 选择缓存目录（传入优先，其次全局默认）
    cache_root = cache_dir or os.path.join(os.path.expanduser("~"), "mineru", "workflow", "search_cache")
    os.makedirs(cache_root, exist_ok=True)

    # Fix: 使用稳定哈希避免重启失效
    cache_file = os.path.join(cache_root, f"{hashlib.md5(query.encode('utf-8')).hexdigest()}.json")
    if os.path.exists(cache_file):
        try:
            with open(cache_file, 'r', encoding='utf-8') as f:
                cached = json.load(f)
            if time.time() - cached.get('timestamp', 0) < 86400:
                print("   💾 [缓存] 使用已缓存搜索结果")
                return cached.get('result', '')
        except Exception as e:
            print(f"   ⚠️ 缓存读取失败，重新搜索: {e}")

    # 非强制时的简单兜底逻辑（保持旧行为）
    if not force:
        urgent_keywords = ['数据', '统计', '最新', '2024', '2025', '报告', '指数', '排名', '分析', '现状']
        if not any(kw in query for kw in urgent_keywords):
            return ""

    proxies = getattr(CONF, "PROXIES_CLOUD", None) or getattr(CONF, "PROXIES_LOCAL", None)
    url = "https://api.tavily.com/search"
    payload = {
        "api_key": api_key,
        "query": query,
        "search_depth": "advanced",
        "include_answer": True,
        "max_results": 5
    }

    for attempt in range(max_retries):
        try:
            print(f"   🌐 [联网搜索] 查询: {query} (尝试 {attempt+1}/{max_retries})")
            res = requests.post(url, json=payload, headers={"Content-Type": "application/json"}, proxies=proxies, timeout=20)
            if res.status_code != 200:
                if res.status_code == 403:
                    print("   ❌ Tavily Key 无效或额度耗尽")
                    return ""
                raise ConnectionError(f"Status {res.status_code}: {res.text[:200]}")

            data = res.json()
            answer = data.get("answer", "")
            details = [r.get("content", "") for r in data.get("results", [])]
            combined = "Answer: " + answer + "\n\nDetails:\n" + "\n".join(details)

            try:
                with open(cache_file, 'w', encoding='utf-8') as f:
                    json.dump({"query": query, "result": combined, "timestamp": time.time()}, f, ensure_ascii=False)
            except Exception:
                pass

            return combined
        except Exception as e:
            if attempt < max_retries - 1:
                wait = 2 ** attempt
                print(f"   ⚠️ 搜索失败 ({e})，{wait}s 后重试...")
                time.sleep(wait)
            else:
                print(f"   ❌ 搜索彻底失败: {e}")
                return ""

# ================== 🛠️ 通用请求 (带重试) ==================

def call_api_robust(url, payload, headers=None, max_retries=3, proxies=None, timeout=600):
    """带重试的 POST 请求包装，返回解析后的 JSON 或 None"""
    headers = headers or {}
    backoff = 2
    for attempt in range(max_retries):
        try:
            res = requests.post(url, json=payload, headers=headers, proxies=proxies, timeout=timeout)
            if res.status_code == 200:
                return res.json()
            if res.status_code in (429, 500, 502, 503, 504):
                print(f"   ⚠️ API 繁忙/异常 ({res.status_code})，{backoff}s 后重试...")
                time.sleep(backoff)
                backoff *= 2
                continue
            print(f"   ❌ API 错误 ({res.status_code}): {res.text[:200]}")
            return None
        except Exception as e:
            print(f"   ⚠️ 连接异常: {e}")
            time.sleep(backoff)
            backoff *= 2
    return None

# ================== 🧠 质量评估与改进 ==================

def validate_json_chart_data(chart_data):
    """
    【数据验证引擎】 检查 JSON 图表数据的完整性和一致性
    返回 (is_valid, error_messages)
    """
    errors = []
    
    # 1. 基础字段检查
    if 'chart_type' not in chart_data:
        errors.append("缺少 chart_type 字段")
    elif chart_data['chart_type'] not in ['bar', 'line', 'pie', 'radar', 'mixed']:
        errors.append(f"无效的 chart_type: {chart_data['chart_type']}")
    
    if 'title' not in chart_data or not chart_data['title']:
        errors.append("缺少 title 字段")
    
    if 'data' not in chart_data:
        errors.append("缺少 data 字段")
        return False, errors
    
    data = chart_data['data']
    
    # 2. 数据结构检查
    if 'labels' not in data or not data['labels']:
        errors.append("缺少或空的 labels")
    
    if 'datasets' not in data or not data['datasets']:
        errors.append("缺少或空的 datasets")
        return False, errors
    
    labels = data.get('labels', [])
    datasets = data.get('datasets', [])
    
    # 3. 数据一致性检查
    for i, ds in enumerate(datasets):
        if 'values' not in ds or not ds['values']:
            errors.append(f"datasets[{i}] 缺少 values 或为空")
            continue
        
        if len(ds['values']) != len(labels):
            errors.append(f"datasets[{i}] 的值个数({len(ds['values'])}) != labels个数({len(labels)})")
        
        # 检查数值类型
        for j, v in enumerate(ds['values']):
            try:
                float(v)
            except (TypeError, ValueError):
                errors.append(f"datasets[{i}].values[{j}] 不是有效数字: {v}")
    
    # 4. 数据范围检查
    if errors:
        return False, errors
    
    # 如果所有检查通过，返回 True
    return True, []


def validate_content_structure(content):
    """
    【内容结构验证】 检查 Markdown 内容的结构完整性
    返回 (is_valid, suggestions)
    """
    suggestions = []
    
    # 1. 检查标题结构
    h3_headers = content.count('###')
    h4_headers = content.count('####')
    
    if h3_headers < 1:
        suggestions.append("建议添加 ### 三级标题来组织内容")
    
    # 2. 检查表格
    table_lines = [l for l in content.split('\n') if '|' in l and l.count('|') >= 3]
    if len(table_lines) < 2:
        suggestions.append("建议添加 Markdown 表格汇总数据")
    
    # 3. 检查JSON图表
    if '```json' not in content or '"chart_type"' not in content:
        suggestions.append("建议添加 JSON 图表数据块")
    
    # 4. 检查数据来源
    if '[' not in content and 'http' not in content:
        suggestions.append("建议为关键数据标注来源")
    
    is_valid = len(suggestions) == 0
    return is_valid, suggestions


# ================== 📝 数据驱动的写作增强 ===================

def extract_data_points_from_content(content):
    """
    从内容中提取数据点和来源
    返回 [(数据, 来源), ...]
    """
    import re
    
    data_points = []
    
    # 提取数字 + 单位（如 "23.5%", "¥100万", "2024年"）
    patterns = [
        r'(\d+(?:\.\d+)?%)',  # 百分比
        r'([¥$€]\s*\d+(?:\.\d+)?(?:万|亿|千)?)',  # 货币
        r'(\d{4}年)',  # 年份
        r'(CAGR\s*[\d.]+%)',  # CAGR
        r'(增长\s*[\d.]+%)',  # 增长率
    ]
    
    for pattern in patterns:
        matches = re.finditer(pattern, content)
        for match in matches:
            # 尝试找到来源标注 [xxx]
            start = max(0, match.start() - 100)
            context = content[start:match.end() + 50]
            source_match = re.search(r'\[([^\]]+)\]', context)
            source = source_match.group(1) if source_match else "未标注"
            data_points.append((match.group(1), source))
    
    return data_points


def evaluate_content_quality(content, topic, section_title):
    """
    【企业级评估系统 v2.0】 返回 (score, feedback, improvement_hints)
    标准：论文级别（逻辑严密、证据充分、结构完善）
    评分体系：0-10分，细粒度反馈
    """
    if not content or len(content) < 300:
        return 0, "内容过短", ["请生成至少300字的内容"]
    
    issues = []
    score = 10.0
    improvement_hints = []
    
    # ========== 1. 内容深度与篇幅 (20%) ==========
    word_count = len(content)
    if word_count < 800:
        issues.append(f"篇幅浅薄 ({word_count}字 < 800字)")
        score -= 2.0
        improvement_hints.append("需要展开论述，至少补充400字以上内容")
    elif word_count < 1200:
        issues.append(f"篇幅中等 ({word_count}字)")
        score -= 0.5
        improvement_hints.append("建议补充案例或对标分析，达到1200字以上")
    
    # ========== 2. 逻辑结构与层级 (20%) ==========
    h3_count = content.count("###")
    h4_count = content.count("####")
    total_headers = h3_count + h4_count
    
    if total_headers == 0:
        issues.append("无逻辑分层")
        score -= 2.5
        improvement_hints.append("添加### 和 #### 标题进行逻辑分层（至少2-3层）")
    elif total_headers == 1:
        issues.append("逻辑层级单薄")
        score -= 1.0
        improvement_hints.append("补充更多子标题，形成树状结构")
    
    # 检查段落长度（过长段落影响可读性）
    paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
    avg_para_len = word_count / max(len(paragraphs), 1)
    if avg_para_len > 400:
        issues.append(f"段落过长 (平均{avg_para_len:.0f}字)")
        score -= 0.8
        improvement_hints.append("建议拆分长段落，每段保持200-300字")
    
    # ========== 3. 数据证据 (25%) ==========
    digit_count = sum(c.isdigit() for c in content)
    data_ratio = digit_count / len(content) if len(content) > 0 else 0
    
    if data_ratio < 0.02:  # 提高到2%
        issues.append(f"数据密度过低 ({data_ratio*100:.1f}% < 2%)")
        score -= 2.5
        improvement_hints.append("补充具体数字、百分比、参数等定量数据（每300字至少1个数据点）")
    elif data_ratio < 0.03:
        score -= 0.5
        improvement_hints.append("数据密度可进一步提升，目标3-5%")
    
    # 检查引用格式（数据来源）
    has_citation = "[" in content and "]" in content
    has_url = "http" in content
    if not has_citation and not has_url:
        issues.append("缺少数据来源标注")
        score -= 1.0
        improvement_hints.append("为关键数据标注来源 [来源] 或 URL")
    
    # ========== 4. 可视化与表格 (20%) ==========
    # 更严格的表格检测
    table_lines = [l for l in content.split('\n') if '|' in l and len(l.split('|')) >= 4]
    has_quality_table = len(table_lines) >= 2
    has_json_chart = "```json" in content and "chart_type" in content
    
    if not has_quality_table and not has_json_chart:
        issues.append("严重缺失可视化")
        score -= 3.0
        improvement_hints.append("必须包含 Markdown 表格或 JSON 图表数据块")
    elif not has_json_chart:
        score -= 1.5
        improvement_hints.append("建议补充 JSON 图表块便于生成专业图表")
    
    # ========== 5. 专业术语密度 (15%) ==========
    professional_keywords = {
        "技术": ["算法", "架构", "模型", "框架", "协议", "接口", "参数", "优化"],
        "商业": ["市场", "竞争", "成本", "效益", "收益", "风险", "战略", "方案"],
        "学术": ["研究", "分析", "论证", "实证", "理论", "假设", "结论"],
        "规范": ["标准", "规范", "符合", "GB", "ISO", "规定", "要求"]
    }
    
    prof_count = 0
    for category, keywords in professional_keywords.items():
        prof_count += sum(content.count(kw) for kw in keywords)
    
    prof_ratio = prof_count / max(word_count / 100, 1)  # 每100字期望1个专业词汇
    
    if prof_ratio < 0.5:
        issues.append(f"专业术语不足 ({prof_ratio:.1f}/100字)")
        score -= 1.5
        improvement_hints.append("增加领域特定的专业术语和行业用语")
    elif prof_ratio < 1.0:
        score -= 0.3
        improvement_hints.append("专业术语使用可进一步增强")
    
    # ========== 6. 论点支撑力 (5%) ==========
    # 检查是否有对比、因果、递进等逻辑词
    logic_connectors = ["因此", "所以", "相比", "与之相对", "进而", "由此可见", "根据", "证明", "表明"]
    logic_count = sum(content.count(conn) for conn in logic_connectors)
    
    if logic_count == 0:
        issues.append("缺乏论证逻辑")
        score -= 1.0
        improvement_hints.append("使用逻辑词强化论证：'因此'、'相比'、'因此'等")
    
    # ========== 最终反馈 ==========
    feedback_text = " | ".join(issues) if issues else "✅ 优秀：结构完善、数据充实、表述专业"
    return max(0.0, score), feedback_text, improvement_hints


def generate_refinement_prompt(original_prompt, content, quality_score, feedback, improvement_hints, round_num):
    """
    【企业级改进系统 v2.0】基于质量评估动态生成改进提示词
    支持多轮迭代，每轮聚焦不同维度
    """
    if quality_score >= 8.5:
        return None  # 质量优秀，无需改进
    
    # 根据改进轮次采用不同策略
    if round_num == 1:
        # 第1轮：快速改进 - 专注篇幅和数据密度
        focus = """
        【第1轮改进：快速深化】
        当前评分: {}/10
        问题汇总: {}
        
        改进重点 (优先级顺序):
        1. 篇幅与深度: 
           - 目标: 1000+ 字（每个子节）
           - 方法: 补充详尽的分析段落、案例论证、对比分析
           
        2. 数据密度:
           - 目标: 数据占比 3-5%（具体数字、百分比、参数）
           - 方法: 加入 CAGR、市场份额、技术参数、成本数据等
           - 格式: 在数据后标注 [来源:xxx] 或 [数据年份:xxxx]
           
        3. 可视化:
           - 必须: 一个高质量 Markdown 表格（汇总数据）
           - 必须: 一个 JSON 图表数据块（便于生成专业图）
        """.format(quality_score, feedback)
        
    elif round_num == 2:
        # 第2轮：逻辑优化 - 聚焦结构和论证
        focus = """
        【第2轮改进：逻辑强化】
        当前评分: {}/10
        改进提示: {}
        
        改进重点 (优先级顺序):
        1. 逻辑结构:
           - 添加 ### 和 #### 子标题（至少3层结构）
           - 每个段落不超过300字，段落间用逻辑词连接
           - 使用逻辑词: 因此、相比、与之相对、由此可见、根据、证明等
           
        2. 论证完善:
           - 每个重要论点后跟具体证据（数据、案例、引用）
           - 采用问题-分析-结论的结构
           - 避免空泛表述，所有观点必须有依据
           
        3. 专业术语:
           - 融入领域特有术语，但保持表达通俗易懂
           - 重要概念首次出现时加括号英文解释
        """.format(quality_score, "\n".join(improvement_hints[:2]))
        
    elif round_num >= 3:
        # 第3轮：精细打磨 - 专业度提升
        focus = """
        【第3轮改进：精细打磨】
        当前评分: {}/10
        
        改进重点 (优先级顺序):
        1. 专业审美提升:
           - 表格格式规范：表头加粗、对齐、单位统一
           - 数据表述：避免"大幅增长"，改用"同比增长 23.5%"
           - 中英文混用：技术术语用英文，业务用中文
           
        2. 内容完善:
           - 补充对标案例或行业典范
           - 加入风险识别或局限性分析
           - 提出可行的改进建议或下一步方向
           
        3. 可读性优化:
           - 复杂概念前加"简言之"的浓缩表述
           - 段落间逻辑递进自然，避免生硬
           - 保留原有优质内容，只补充不删减
        """.format(quality_score)
    else:
        focus = "基于反馈补充改进"
    
    refinement_prompt = f"""
    {original_prompt}
    
    ================== 多轮迭代改进 ==================
    {focus}
    
    【必须遵守的格式要求】:
    1. 返回完整的改进后内容（不是修改说明）
    2. 保留原有的好内容，只补充不删减
    3. 保持 Markdown 格式一致
    4. 包含表格时，行数 >= 4（表头+数据行）
    5. 包含 JSON 图表数据时，确保语法正确（无注释）
    """
    return refinement_prompt


def generate_enhanced_prompt_with_visuals(base_prompt, topic, section_title):
    """
    【升级版】生成增强的提示词
    角色：行业领域资深专家 (Industry Domain Senior Expert)
    策略：CoT (思维链) + 深度行业洞察 + 强制 JSON 数据输出
    """
    enhanced = f"""
    {base_prompt}
    
    ========== 🏛️ 角色与标准 ==========
    你现在是【{topic}】领域的【首席行业专家】(Chief Industry Expert)。
    你拥有30年的行业实战经验，对该领域的技术演进、市场格局、政策走向有极深的洞察。
    
    你的任务是撰写一份具有【行业标杆水准】的深度研究报告/技术方案。
    
    【专家级写作标准】：
    1. **深度洞察 (Deep Insight)**：
       - 不要只罗列现象，要挖掘背后的根本原因 (Root Cause) 和底层逻辑。
       - 能够识别行业痛点，并提出切实可行的、具有前瞻性的解决方案。
       
    2. **技术权威 (Technical Authority)**：
       - 熟练引用最新的行业标准 (GB/ISO/IEC)、专利技术或学术前沿成果。
       - 使用该领域最地道的行话 (Jargon)，但要解释清楚其商业价值。
       
    3. **数据驱动 (Data-Driven)**：
       - 每一个论点都必须有具体的数据支撑（如：CAGR、市场渗透率、技术参数、成本效益比）。
       - 拒绝模糊的“大幅提升”，必须量化为“提升了 23.5%”。
       
    4. **结构严谨 (Structured Thinking)**：
       - 使用 MECE 原则（相互独立，完全穷尽）组织内容。
       - 采用“结论先行”的咨询顾问式表达风格。
    
    ========== 📊 强制可视化要求 (JSON) ==========
    为了生成专业图表，你必须在回答末尾提供一个【标准的 JSON 数据块】，格式如下：
    
    ```json
    {{
      "chart_type": "bar",  // 可选: "bar"(柱状), "line"(折线), "pie"(饼图), "radar"(雷达图), "mixed"(组合)
      "title": "2020-2024年市场份额分析",
      "data": {{
        "labels": ["2020", "2021", "2022", "2023", "2024"],
        "datasets": [
           {{ "label": "本公司", "values": [20, 25, 30, 35, 40] }},
           {{ "label": "竞争对手A", "values": [30, 28, 25, 20, 15] }}
        ]
      }},
      "x_label": "年份",
      "y_label": "市场份额 (%)"
    }}
    ```
    *注意：请确保 JSON 语法完全正确，不要在 JSON 中添加注释。*
    
    ========== 📝 专家视角结构建议 ==========
    1. **核心论点 (Thesis Statement)**: 开篇即亮出具有专家视角的独到见解。
    2. **现状与挑战 (Status & Challenges)**: 深度剖析当前的技术/市场瓶颈。
    3. **解决方案 (Solution)**: 提出系统性的技术路线或战略建议，强调创新点。
    4. **价值验证 (Value Proposition)**: 用数据和案例证明方案的可行性与优越性。
    
    关键词：{topic}, {section_title}
    """
    return enhanced

def generate_enhanced_prompt_with_visuals(base_prompt, topic, section_title, attempt_num=1):
    """
    【企业级提示词工程 v3.0】 
    采用高级策略：角色扮演 + CoT (思维链) + 行业标杆 + 强制输出格式
    attempt_num: 第几轮尝试（1=初稿，2+=改进）
    """
    
    # ========== 动态调整策略 ==========
    if attempt_num == 1:
        # 初稿：强调深度和专业度
        thinking_depth = "深度反思"
        output_target = "行业标杆级"
    else:
        # 改进稿：强调数据和严谨
        thinking_depth = "详尽论证"
        output_target = "专业期刊级"
    
    enhanced = f"""
    【前置任务分析】
    你即将生成一份关于 【{topic}】 中 【{section_title}】 的专业内容。
    这份内容将作为行业研究报告的重要章节，目标受众是行业高管、技术决策者和投资方。
    质量标准：{output_target}水平。
    
    ========== 🏛️ 角色定位 ==========
    你是该领域的【首席专业咨询师】(Chief Professional Consultant)
    • 拥有15+年的行业深耕经历
    • 曾撰写多份行业标杆研究报告
    • 对市场、技术、政策有系统性认知
    • 擅长用数据和案例支撑论点
    
    ========== 🧠 {thinking_depth}流程 (CoT - Chain of Thought) ==========
    在开始写作前，请按以下步骤{thinking_depth}：
    
    1. 【背景梳理】
       - 这个主题在行业中的当前状态？
       - 存在哪些核心问题和痛点？
       - 为什么这个话题对目标受众重要？
    
    2. 【论点构建】
       - 核心观点是什么（一句话总结）？
       - 有哪些支撑这个观点的证据（数据、案例、理论）？
       - 是否存在对立观点？如何驳斥或融合？
    
    3. 【结构规划】
       - 如何分层组织内容（3-4层逻辑结构）？
       - 段落顺序是否形成自然递进？
       - 是否需要可视化补充？
    
    4. 【质量检查】
       - 数据密度是否充足（3-5%）？
       - 专业术语使用是否得当？
       - 逻辑连接词是否充分？
    
    ========== 📝 内容写作标准 ==========
    
    【深度洞察 (Deep Insight)】
    • 不仅描述现象，更要分析根本原因和驱动因素
    • 识别隐含的行业痛点和机遇
    • 提出前瞻性、可行性的解决方案
    
    【数据驱动 (Data-Driven)】
    • 每个重要论点后必须有具体数据支撑
    • 数据格式：具体数字、百分比、CAGR、同比/环比增长
    • 示例：
      ❌ "市场增长很快"
      ✅ "2023年市场规模为 ¥XX 亿，CAGR 为 23.5%，预计2025年达 ¥YY 亿"
    • 数据标注：在数据后加 [来源:XXX 2024年报] 或 [数据截至2024年10月]
    
    【技术权威 (Technical Authority)】
    • 融入行业标准、国际规范、技术框架
    • 使用行业特有术语，但保持可理解
    • 引用如 GB、ISO、IEC、IEEE、RFC 等
    • 第一次出现技术术语时，可用括号补充英文
    
    【结构严谨 (Structured Thinking)】
    • 采用MECE原则（相互独立，完全穷尽）组织论点
    • 使用"结论先行"风格：先说结论，再说原因
    • 段落间逻辑词连接：因此、相比、与之相对、进而、由此可见等
    • Markdown 结构：
      # 一级标题（章名）
      ## 二级标题（节名）
      ### 三级标题（核心论点）
      #### 四级标题（分论点）
    
    【可视化丰富 (Visual Rich)】
    • Markdown 表格：汇总关键数据或对比分析
      - 行数 >= 4（表头+数据行）
      - 列数 >= 3（明确的对比维度）
      - 表格前需有解释文字
    
    • JSON 图表数据块：用于生成专业图表
      - 放在正文末尾，```json ... ``` 格式
      - 必须包含：chart_type, title, data (labels, datasets), x_label, y_label
      - 【重要】尽量使用多样化图表类型，不要重复：
        * bar - 柱状图（对比、排序、排名）
        * line - 折线图（趋势、时序、增长率）
        * pie - 饼图（占比、市场份额、构成）
        * radar - 雷达图（多维评价、技术对标、能力评分）
        * area - 面积图（堆积趋势、成本分解）
        * scatter - 散点图（相关性、分布）
        * bubble - 气泡图（三维对比、体量差异）
        * stacked_bar - 堆积柱状图（结构分析、层级分解）
        * heatmap - 热力图（矩阵、密度、产业链）
        * mixed - 混合图（多指标、复合分析）
      - 【禁止】不要在同一份报告中重复使用相同图表类型
    
    ========== 🎯 具体写作要求 ==========
    
    1. 【篇幅与质量】
       - 目标字数：1000-1500 字
       - 段落数：5-8 段
       - 每段长度：150-300 字
       - 每段前有逻辑起点（主题句）
    
    2. 【数据与案例】
       - 至少包含 5-10 个数据点（具体数字）
       - 至少包含 2-3 个真实案例或对标分析
       - 数据密度目标：每300字至少1个数据点
    
    3. 【表格与图表 - 多样化】
       - 必须包含：1个 Markdown 表格（4行以上）
       - 必须包含：1个 JSON 图表块（选择最合适的图表类型）
       - 【关键】每个子节选择不同的图表类型，保持视觉多样性
       - 必须包含：1个 JSON 图表数据块（完整的 chart_type/data 结构）
       - 表格和图表前需有过渡文字
    
    4. 【术语与表述】
       - 专业术语占比：每100字至少1个行业术语
       - 避免模糊表述：用具体量词代替"很"、"比较"、"大幅"
       - 示例：
         ❌ "成本比较高"
         ✅ "单位成本为 ¥X/件，相比竞品高 23%"
    
    5. 【逻辑与论证】
       - 使用逻辑词强化论证：因此、相比、基于、证明、表明、进而、由此可见等
       - 避免生硬的列举，应该形成因果链
       - 每个论点后跟证据（如果未提供证据，则说明需要补充）
    
    ========== 📋 JSON 图表示例 ==========
    
    ```json
    {{
      "chart_type": "line",
      "title": "2020-2024年行业市场规模与增长率",
      "data": {{
        "labels": ["2020", "2021", "2022", "2023", "2024E"],
        "datasets": [
          {{"label": "市场规模(¥亿)", "values": [100, 130, 169, 220, 286]}},
          {{"label": "YoY增长率(%)", "values": [15, 30, 30, 30, 30]}}
        ]
      }},
      "x_label": "年份",
      "y_label": "规模/增长率"
    }}
    ```
    
    ========== 📊 支持的图表类型（9种） ==========
    
    1. **bar** - 柱状图（对比、排序）
    2. **line** - 折线图（趋势、时序）
    3. **pie** - 饼图（占比、分布）
    4. **radar** - 雷达图（多维评价、对标）
    5. **area** - 面积图（堆积、趋势）
    6. **scatter** - 散点图（相关性、离散度）
    7. **bubble** - 气泡图（三维对比）
    8. **stacked_bar** - 堆积柱状图（结构分析）
    9. **heatmap** - 热力图（矩阵、密度）
    10. **mixed** - 混合图（柱+折线）
    
    ========== 💡 选型建议 ==========
    - 数据对标、市场排名 → 柱状图 (bar)
    - 市场趋势、增长率 → 折线图 (line)
    - 市场份额、占比 → 饼图 (pie)
    - 技术评分、多维评价 → 雷达图 (radar)
    - 成本构成、层级分解 → 堆积柱状图 (stacked_bar)
    - 产业链分析、性能对比 → 热力图 (heatmap)
    - 相关性分析、散布情况 → 散点图 (scatter)
    - 快速增长、体量对比 → 气泡图 (bubble)
    - 多个指标混合 → 混合图 (mixed)
    - 面积堆积分析 → 面积图 (area)
    
    ========== 🔍 最后检查清单 ==========
    在提交内容前，自检：
    □ 篇幅 >= 1000 字？
    □ 有3-4层逻辑结构（###/#### 标题）？
    □ 数据密度 >= 3%？
    □ 至少5个数据点，标注了来源？
    □ 包含1个表格（4行以上）？
    □ 包含1个完整的JSON图表块？
    □ 使用了5个以上逻辑词？
    □ 专业术语自然融入？
    □ 段落间逻辑递进清晰？
    
    ========== 实际写作内容 ==========
    {base_prompt}
    """
    return enhanced

def create_chart_from_description_plotly(chart_data, output_path):
    """
    【企业级图表 v2.0】使用 Plotly 生成交互式专业图表
    支持: Bar, Line, Pie, Radar, Mixed
    输出: HTML（可交互）和 PNG（静态）
    """
    try:
        import plotly.graph_objects as go
        import plotly.io as pio
        
        # 数据验证
        if not chart_data or not isinstance(chart_data, dict):
            return False
        
        c_type = chart_data.get('chart_type', 'bar').lower()
        title = chart_data.get('title', '数据分析')
        data = chart_data.get('data', {})
        labels = data.get('labels', [])
        datasets = data.get('datasets', [])
        x_label = chart_data.get('x_label', '')
        y_label = chart_data.get('y_label', '')
        
        # 基础验证
        if not labels or not datasets:
            return False
        
        fig = go.Figure()
        
        # ========== 柱状图、折线图、混合图 ==========
        if c_type in ['bar', 'line', 'mixed']:
            for idx, ds in enumerate(datasets):
                if not isinstance(ds, dict):
                    continue
                label = ds.get('label', f'Series {idx+1}')
                vals = ds.get('values', [])
                
                if not vals or len(vals) != len(labels):
                    continue
                
                if c_type == 'line' or (c_type == 'mixed' and idx > 0):
                    fig.add_trace(go.Scatter(
                        x=labels, y=vals, mode='lines+markers',
                        name=label, line=dict(width=3),
                        hovertemplate='<b>%{x}</b><br>' + label + ': %{y}<extra></extra>'
                    ))
                else:
                    fig.add_trace(go.Bar(
                        x=labels, y=vals, name=label,
                        hovertemplate='<b>%{x}</b><br>' + label + ': %{y}<extra></extra>'
                    ))
        
        # ========== 饼图 ==========
        elif c_type == 'pie':
            if datasets and isinstance(datasets[0], dict) and 'values' in datasets[0]:
                values = datasets[0]['values']
                if values and len(values) == len(labels):
                    fig = go.Figure(data=[go.Pie(
                        labels=labels, values=values,
                        hovertemplate='<b>%{label}</b><br>占比: %{value}<extra></extra>'
                    )])
                else:
                    return False
            else:
                return False
        
        # ========== 雷达图 ==========
        elif c_type == 'radar':
            for ds in datasets:
                if not isinstance(ds, dict):
                    continue
                label = ds.get('label', '')
                vals = ds.get('values', [])
                if vals and len(vals) == len(labels):
                    fig.add_trace(go.Scatterpolar(
                        r=vals, theta=labels, fill='toself',
                        name=label,
                        hovertemplate='<b>%{theta}</b><br>' + label + ': %{r}<extra></extra>'
                    ))
        
        # ========== 布局优化 ==========
        fig.update_layout(
            title={'text': title, 'font': {'size': 18, 'color': '#1f77b4'}},
            xaxis_title=x_label,
            yaxis_title=y_label,
            hovermode='x unified',
            template='plotly_white',
            font=dict(family="Arial, sans-serif", size=12),
            plot_bgcolor='rgba(240, 240, 240, 0.5)',
            paper_bgcolor='white',
            width=1000,
            height=600,
            showlegend=True,
            legend=dict(yanchor="top", y=0.99, xanchor="left", x=0.01)
        )
        
        # ========== 保存输出 ==========
        html_path = output_path.replace('.png', '.html')
        fig.write_html(html_path)
        print(f"      🖼️ Plotly HTML 已生成: {html_path}")
        
        # 同时保存PNG (静态版本) - 可选
        try:
            pio.write_image(fig, output_path, width=1000, height=600)
            print(f"      📊 Plotly PNG 已生成: {output_path}")
            return True
        except Exception as png_err:
            # Kaleido可能未安装，只保存HTML也可以
            print(f"      ℹ️ PNG 输出需要 kaleido 库 (已保存 HTML 版本)")
            return True
            
    except ImportError as ie:
        print(f"      ℹ️ Plotly 不可用: {ie}")
        return False
    except Exception as e:
        print(f"      ⚠️ Plotly 处理异常: {type(e).__name__}: {str(e)[:80]}")
        return False

def create_chart_from_description(chart_desc, output_path):
    """
    【重构版】从 JSON 数据生成商业级 (Business-Level) 图表
    支持类型: Bar, Line, Pie, Radar, Mixed
    风格: 商务蓝调, 高分辨率, 中文支持
    【新增】数据验证和空白图表检测
    """
    try:
        # 1. 提取 JSON 数据
        json_match = re.search(r'```json\s*(\{.*?\})\s*```', chart_desc, re.DOTALL)
        if not json_match:
            # 尝试直接找 {}
            json_match = re.search(r'(\{.*"chart_type".*\})', chart_desc, re.DOTALL)
            
        if not json_match:
            print("      ℹ️ 未检测到有效 JSON 图表数据，跳过绘图")
            return False
            
        data_str = json_match.group(1)
        try:
            chart_data = json.loads(data_str)
        except json.JSONDecodeError as je:
            print(f"      ⚠️ JSON 解析失败: {str(je)[:80]}")
            return False

        # 【新增】数据验证
        labels = chart_data.get('data', {}).get('labels', [])
        datasets = chart_data.get('data', {}).get('datasets', [])
        
        # 检查是否为空白或无效数据
        if not labels or not datasets:
            print("      ⚠️ 图表数据为空，跳过绘图")
            return False
        
        # 检查数据一致性
        for ds in datasets:
            if not ds.get('values') or len(ds['values']) != len(labels):
                print(f"      ⚠️ 数据不一致：values长度 {len(ds.get('values', []))} != labels长度 {len(labels)}")
                return False
        
        # 检查数据是否全为0或无效
        all_values = []
        for ds in datasets:
            all_values.extend(ds.get('values', []))
        if not all_values or all(v == 0 or v is None for v in all_values):
            print("      ⚠️ 所有数据点均为0或空，生成备选表格")
            return False

        # ========== 优先尝试 Plotly（交互式） ==========
        if create_chart_from_description_plotly(chart_data, output_path):
            return True
        
        # ========== 降级到 Matplotlib（静态） ==========
        print("      📊 使用 Matplotlib 生成静态图表...")
        
        # 2. 准备绘图上下文
        plt.clf()
        plt.close('all')
        
        # 设置商业风格 (Business Style)
        # 背景色: 浅灰/白, 网格: 灰色虚线, 字体色: 深灰
        plt.rcParams.update({
            'figure.facecolor': '#FFFFFF',
            'axes.facecolor': '#F8F9FA',
            'axes.edgecolor': '#DEE2E6',
            'axes.grid': True,
            'grid.color': '#E9ECEF',
            'grid.linestyle': '--',
            'grid.alpha': 0.8,
            'text.color': '#343A40',
            'axes.labelcolor': '#495057',
            'xtick.color': '#495057',
            'ytick.color': '#495057',
            'font.size': 10
        })
        
        # 商业配色方案 (深蓝, 科技蓝, 活力橙, 稳重灰)
        COLORS = ['#0056B3', '#20C997', '#FD7E14', '#6C757D', '#6610F2', '#E83E8C']
        
        fig, ax = plt.subplots(figsize=(10, 6), dpi=150) # 提高分辨率
        
        # 字体应用 (确保中文)
        font_prop = CHINESE_FONT if CHINESE_FONT else None
        
        # 3. 解析数据
        c_type = chart_data.get('chart_type', 'bar').lower()
        title = chart_data.get('title', '数据分析')
        x_label = chart_data.get('x_label', '')
        y_label = chart_data.get('y_label', '')

        # 4. 绘制逻辑 - 支持8+种图表类型
        if c_type == 'pie':
            # 【饼图】- 占比分析
            if not datasets or 'values' not in datasets[0]:
                return False
            values = datasets[0]['values']
            pie_labels = labels
            textprops_dict = {}
            if font_prop:
                textprops_dict['fontproperties'] = font_prop
            wedges, texts, autotexts = ax.pie(
                values, labels=pie_labels, autopct='%1.1f%%',
                startangle=90, colors=COLORS,
                textprops=textprops_dict
            )
            for text in autotexts: 
                text.set_color('white')

        elif c_type == 'radar':
            # 【雷达图 v2.0】- 多维度对比 (已修复显示问题)
            plt.close(fig)
            fig = plt.figure(figsize=(10, 8), dpi=150)
            ax = fig.add_subplot(111, polar=True)
            
            # 使用 Python 内置替代 numpy
            import math
            num_vars = len(labels)
            angles = [2 * math.pi * i / num_vars for i in range(num_vars)]
            angles += angles[:1]  # 闭合
            
            # 设置雷达图标签
            ax.set_xticks(angles[:-1])
            ax.set_xticklabels(labels, fontproperties=font_prop, size=10)
            ax.set_ylim(0, 100)  # 设置半径范围
            ax.set_rlabel_position(0)  # 标签位置
            
            # 绘制每个数据集
            for idx, ds in enumerate(datasets):
                vals = ds.get('values', [])
                if len(vals) != num_vars:
                    continue
                vals_plot = vals + vals[:1]  # 闭合
                ax.plot(angles, vals_plot, 'o-', linewidth=2.5, label=ds.get('label', f'数据{idx+1}'),
                        color=COLORS[idx % len(COLORS)])
                ax.fill(angles, vals_plot, alpha=0.25, color=COLORS[idx % len(COLORS)])
            
            ax.legend(loc='upper right', bbox_to_anchor=(1.3, 1.1), fontsize=10)
            ax.grid(True, linestyle='--', alpha=0.7)

        elif c_type == 'area':
            # 【面积图】- 趋势堆积分析
            x = range(len(labels))
            bottom = None
            for idx, ds in enumerate(datasets):
                vals = ds.get('values', [])
                if len(vals) == len(labels):
                    ax.fill_between(x, 0 if bottom is None else bottom, 
                                   [v + (0 if bottom is None else b) for v, b in zip(vals, bottom or [0]*len(vals))],
                                   label=ds.get('label', f'Series {idx+1}'),
                                   color=COLORS[idx % len(COLORS)], alpha=0.7)
                    bottom = [v + (b or 0) for v, b in zip(vals, bottom or [0]*len(vals))]
            
            ax.set_xticks(x)
            ax.set_xticklabels(labels, fontproperties=font_prop)
            if x_label: ax.set_xlabel(x_label, fontproperties=font_prop)
            if y_label: ax.set_ylabel(y_label, fontproperties=font_prop)
            ax.legend(prop=font_prop)

        elif c_type == 'scatter':
            # 【散点图】- 相关性分析
            for idx, ds in enumerate(datasets):
                vals = ds.get('values', [])
                if len(vals) == len(labels):
                    x_vals = range(len(labels))
                    ax.scatter(x_vals, vals, s=100, alpha=0.6, 
                              label=ds.get('label', f'Series {idx+1}'),
                              color=COLORS[idx % len(COLORS)])
            
            ax.set_xticks(range(len(labels)))
            ax.set_xticklabels(labels, fontproperties=font_prop)
            if x_label: ax.set_xlabel(x_label, fontproperties=font_prop)
            if y_label: ax.set_ylabel(y_label, fontproperties=font_prop)
            ax.legend(prop=font_prop)

        elif c_type == 'stacked_bar':
            # 【堆积柱状图】- 组成结构分析
            x = range(len(labels))
            bottom = [0] * len(labels)
            for idx, ds in enumerate(datasets):
                vals = ds.get('values', [])
                if len(vals) == len(labels):
                    ax.bar(x, vals, bottom=bottom, label=ds.get('label', f'Series {idx+1}'),
                          color=COLORS[idx % len(COLORS)], alpha=0.9)
                    bottom = [b + v for b, v in zip(bottom, vals)]
            
            ax.set_xticks(x)
            ax.set_xticklabels(labels, fontproperties=font_prop)
            if x_label: ax.set_xlabel(x_label, fontproperties=font_prop)
            if y_label: ax.set_ylabel(y_label, fontproperties=font_prop)
            ax.legend(prop=font_prop)

        elif c_type == 'bubble':
            # 【气泡图】- 三维数据对比
            for idx, ds in enumerate(datasets):
                vals = ds.get('values', [])
                if len(vals) == len(labels):
                    x_vals = range(len(labels))
                    sizes = [abs(v) * 10 + 50 for v in vals]  # 气泡大小
                    ax.scatter(x_vals, vals, s=sizes, alpha=0.5,
                              label=ds.get('label', f'Series {idx+1}'),
                              color=COLORS[idx % len(COLORS)])
            
            ax.set_xticks(range(len(labels)))
            ax.set_xticklabels(labels, fontproperties=font_prop)
            if x_label: ax.set_xlabel(x_label, fontproperties=font_prop)
            if y_label: ax.set_ylabel(y_label, fontproperties=font_prop)
            ax.legend(prop=font_prop)

        elif c_type == 'heatmap':
            # 【热力图】- 矩阵数据分析（需要 numpy 库支持）
            try:
                import numpy as np
                data_matrix = []
                for ds in datasets:
                    vals = ds.get('values', [])
                    if len(vals) == len(labels):
                        data_matrix.append(vals)
                
                if data_matrix:
                    data_matrix = np.array(data_matrix)
                    im = ax.imshow(data_matrix, cmap='RdYlBu_r', aspect='auto')
                    ax.set_xticks(range(len(labels)))
                    ax.set_yticks(range(len(datasets)))
                    ax.set_xticklabels(labels, fontproperties=font_prop)
                    ax.set_yticklabels([ds.get('label', f'Row {i}') for i, ds in enumerate(datasets)], 
                                      fontproperties=font_prop)
                    # 添加颜色条
                    cbar = plt.colorbar(im, ax=ax)
                    cbar.set_label(y_label or '数值', fontproperties=font_prop)
            except ImportError:
                # numpy 不可用，显示文本提示
                print(f"      ℹ️ 热力图需要 numpy 库支持，已跳过此图表类型")
                ax.text(0.5, 0.5, 'Heatmap 图表类型\n需要 numpy 库支持', 
                       ha='center', va='center', transform=ax.transAxes, 
                       fontproperties=font_prop, color='gray', fontsize=12)
                ax.set_xticks([])
                ax.set_yticks([])

        else:
            # Bar / Line / Mixed (默认)
            x = range(len(labels))
            width = 0.35
            
            for idx, ds in enumerate(datasets):
                label = ds.get('label', f'Series {idx+1}')
                vals = ds.get('values', [])
                
                if len(vals) != len(labels): 
                    continue
                
                current_type = c_type
                if c_type == 'mixed':
                    current_type = 'bar' if idx == 0 else 'line'
                
                if current_type == 'line':
                    ax.plot(labels, vals, marker='o', linewidth=2.5,
                            color=COLORS[idx % len(COLORS)], label=label, markersize=8)
                else:
                    offset = (idx - len(datasets)/2) * width + width/2
                    if c_type == 'mixed': offset = 0
                    
                    rects = ax.bar([i + offset for i in x], vals, width,
                           label=label, color=COLORS[idx % len(COLORS)], alpha=0.9)
                    ax.bar_label(rects, padding=3, fmt='%.1f', fontsize=8)

            ax.set_xticks(x)
            ax.set_xticklabels(labels, fontproperties=font_prop)
            if x_label: ax.set_xlabel(x_label, fontproperties=font_prop)
            if y_label: ax.set_ylabel(y_label, fontproperties=font_prop)
            ax.legend(prop=font_prop)

        # 5. 通用修饰
        ax.set_title(title, fontproperties=font_prop, fontsize=16, fontweight='bold', pad=25, color='#212529')
        
        # 移除顶部和右侧边框 (Clean Look)
        if c_type != 'heatmap':
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
        
        plt.tight_layout()
        plt.savefig(output_path, format='png', bbox_inches='tight', dpi=150)
        plt.close()
        
        print(f"      ✅ 图表已生成: {output_path}")
        return True
    except Exception as e:
        print(f"      ⚠️ 图表生成失败: {e}")
        return False


def embed_chart_in_markdown(content, section_title, chapter_title, output_dir):
    """
    检测内容中的图表描述，生成实际图表并嵌入Markdown
    会自动将生成图表的代码块 **替换** 为图片，避免文中残留代码
    【新增】如果图表生成失败，自动降级为表格
    """
    # 检查是否有图表描述标记
    if '```' not in content and '|' not in content:
        return content
    
    safe_title = re.sub(r'[^\w\-]', '_', section_title)[:20]
    charts_dir = os.path.join(output_dir, "charts")
    
    try:
        os.makedirs(charts_dir, exist_ok=True)
        
        # 使用正则查找所有代码块 (包括 python 和 纯文本)
        # 捕获组 1: 完整的代码块 (用于替换)
        # 捕获组 2: 代码块内部的内容 (用于解析)
        pattern = re.compile(r'(```(?:python|py|json)?\s*(\n.*?)```)', re.IGNORECASE | re.DOTALL)
        
        matches = pattern.findall(content)
        
        for i, (full_block, inner_code) in enumerate(matches):
            # 判断是否是图表相关的代码块
            is_chart_code = "matplotlib" in inner_code or "plt." in inner_code
            is_chart_data = "图表" in inner_code and "|" in inner_code
            is_json_chart = '"chart_type"' in inner_code and '"data"' in inner_code
            
            if is_chart_code or is_chart_data or is_json_chart:
                # 生成唯一文件名
                unique_id = f"{int(time.time())}_{i}"
                filename = f"{safe_title}_{unique_id}.png"
                chart_abs_path = os.path.join(charts_dir, filename)
                chart_rel_path = f"./charts/{filename}"
                
                # 尝试生成图表
                success = create_chart_from_description(inner_code, chart_abs_path)
                
                if success and os.path.exists(chart_abs_path):
                    # 图表生成成功
                    img_tag = f"\n\n![{section_title}数据分析]({chart_rel_path})\n\n"
                    content = content.replace(full_block, img_tag)
                    print(f"      ✅ 已生成图表: {filename}")
                else:
                    # 图表生成失败，尝试从JSON中提取表格数据作为备选
                    print(f"      ⚠️ 图表生成失败，尝试备选表格方案...")
                    
                    # 从 JSON 中提取数据生成表格
                    json_match = re.search(r'(\{.*"chart_type".*\})', inner_code, re.DOTALL)
                    if json_match:
                        try:
                            chart_data = json.loads(json_match.group(1))
                            fallback_table = _generate_fallback_table(chart_data)
                            if fallback_table:
                                content = content.replace(full_block, f"\n\n{fallback_table}\n\n")
                                print(f"      ✅ 已生成备选表格")
                                continue
                        except:
                            pass
                    
                    # 如果都失败，保留原代码块但标记为待处理
                    marked_block = f"\n\n⚠️ [图表生成失败 - 请手动处理]\n{full_block}\n\n"
                    content = content.replace(full_block, marked_block)
                    print(f"      ⚠️ 已标记为待处理")
                
    except Exception as e:
        print(f"      ⚠️ 嵌入图表失败: {e}")
    
    return content


def _generate_fallback_table(chart_data):
    """
    从图表JSON数据生成Markdown表格作为备选方案
    """
    try:
        labels = chart_data.get('data', {}).get('labels', [])
        datasets = chart_data.get('data', {}).get('datasets', [])
        title = chart_data.get('title', '数据表')
        
        if not labels or not datasets:
            return None
        
        # 构建表格
        header = '| ' + ' | '.join(['指标'] + labels) + ' |'
        separator = '| ' + ' | '.join(['---'] * (len(labels) + 1)) + ' |'
        
        rows = []
        for ds in datasets:
            label = ds.get('label', '数据')
            values = ds.get('values', [])
            if len(values) == len(labels):
                row = '| ' + ' | '.join([label] + [str(v) for v in values]) + ' |'
                rows.append(row)
        
        if rows:
            table = f"\n\n**表: {title}**\n\n{header}\n{separator}\n" + "\n".join(rows) + "\n\n"
            return table
    except:
        pass
    
    return None


# ================== 📚 增强的 RAG 与上下文管理 ==================

class ContextManager:
    """管理已生成内容的上下文，支持交叉引用"""
    def __init__(self):
        self.generated_sections = {}  # {chapter_title: {section_title: content}}
        self.section_summaries = {}   # 快速索引
        self.cache = {
            "outline": None,
            "style_guide": "",
            "section_plans": {},
            "global_thesis": "",
            "last_exec_summary": ""
        }
    
    def set_master_plan(self, outline, style_guide):
        """缓存 Master Planner 结果（大纲 + 风格指南）"""
        self.cache["outline"] = outline
        self.cache["style_guide"] = style_guide or ""
    
    def set_global_thesis(self, thesis):
        self.cache["global_thesis"] = thesis or ""
    
    def get_global_thesis(self):
        return self.cache.get("global_thesis", "")
    
    def set_last_exec_summary(self, summary):
        self.cache["last_exec_summary"] = summary or ""
    
    def get_last_exec_summary(self):
        return self.cache.get("last_exec_summary", "")

    def add_section_plan(self, section_key, plan_text):
        """缓存 Section Planner 的拆解计划"""
        if plan_text:
            self.cache["section_plans"][section_key] = plan_text
    
    def get_style_guide(self):
        return self.cache.get("style_guide", "")
    
    def get_outline(self):
        return self.cache.get("outline")
    
    def add_section(self, chapter, section, content):
        """添加生成的章节"""
        if chapter not in self.generated_sections:
            self.generated_sections[chapter] = {}
        self.generated_sections[chapter][section] = content
        
        # 生成摘要便于快速引用
        summary = content[:200] + "..." if len(content) > 200 else content
        self.section_summaries[f"{chapter}_{section}"] = summary
    
    def get_related_context(self, topic, max_sections=3):
        """获取相关的已生成内容作为上下文"""
        if not self.generated_sections:
            return ""
        
        # 简单的相关性检索
        related = []
        for chapter, sections in self.generated_sections.items():
            for section, content in sections.items():
                # 检查是否包含关键词
                if any(word in content for word in topic.split()):
                    related.append(f"[{chapter} - {section}]\n{content[:300]}...\n")
        
        return "\n".join(related[:max_sections])
    
    def get_summary(self):
        """生成已生成内容的摘要"""
        summary = "已生成章节：\n"
        for chapter, sections in self.generated_sections.items():
            summary += f"- {chapter}:\n"
            for section in sections.keys():
                summary += f"  * {section}\n"
        return summary

def get_model(task_type):
    """
    极致省钱的模型路由：
    - deep_thinking: 大纲/润色/写代码 → Gemini 3 Pro
    - logic_planning: 拆解章节/简单审核 → Gemini 2.5 Flash（可按质量切换）
    - heavy_reading: 搜索、阅读网页、数据清洗 → Gemini 2.5 Flash Lite
    """
    router = {
        "deep_thinking": CONF.GEMINI_PRO_MODEL,
        "logic_planning": CONF.GEMINI_FLASH_MODEL,
        "heavy_reading": CONF.GEMINI_FLASH_LITE_MODEL
    }
    return router.get(task_type, CONF.GEMINI_PRO_MODEL)

def call_model(prompt, model_id, temperature=0.6, response_mime_type="text/plain"):
    """通用 Gemini 调用封装，支持选择模型/温度"""
    api_version = "v1beta"
    url = f"https://generativelanguage.googleapis.com/{api_version}/models/{model_id}:generateContent?key={CONF.GEMINI_API_KEY}"
    payload = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {
            "temperature": temperature,
            "top_p": 0.95,
            "top_k": 40,
            "response_mime_type": response_mime_type
        }
    }
    resp = call_api_robust(url, payload, headers={"Content-Type": "application/json"}, proxies=CONF.PROXIES_CLOUD, timeout=120)
    if not resp:
        # 允许直连兜底
        resp = call_api_robust(url, payload, headers={"Content-Type": "application/json"}, proxies=CONF.PROXIES_LOCAL, timeout=120)
    if not resp:
        return None
    try:
        return resp['candidates'][0]['content']['parts'][0]['text']
    except Exception:
        print("❌ Gemini 返回结构异常")
        return None

def call_flash(prompt, temperature=0.6, task_type="logic_planning"):
    """Flash 模型（蓝领：搜索/粗写/数据整理）"""
    return call_model(prompt, get_model(task_type), temperature=temperature)

def call_pro(prompt, temperature=0.4, response_mime_type="text/plain"):
    """Pro 模型（白领：规划/润色/写代码）"""
    return call_model(prompt, get_model("deep_thinking"), temperature=temperature, response_mime_type=response_mime_type)

def call_gemini(prompt, json_mode=False):
    """调用 Gemini (用于大纲/统筹)，带明确的错误处理
    使用 gemini-3-pro-preview 获得最优的结构化输出"""
    model_id = CONF.GEMINI_OUTLINE_MODEL
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_id}:generateContent?key={CONF.GEMINI_API_KEY}"
    payload = {
        "contents": [{"parts": [{"text": prompt}]}], 
        "generationConfig": {
            "response_mime_type": "application/json" if json_mode else "text/plain",
            "temperature": 0.3,  # 大纲生成使用更低的温度确保准确性
            "top_p": 0.95,
            "top_k": 40
        }
    }
    
    # 先走代理，失败后直连重试一次；代理不可用时直接直连
    resp = call_api_robust(url, payload, headers={"Content-Type": "application/json"}, proxies=CONF.PROXIES_CLOUD, timeout=120)
    if not resp:
        print("   ↻ 尝试直连 Gemini 再试一次...")
        resp = call_api_robust(url, payload, headers={"Content-Type": "application/json"}, proxies={"http": None, "https": None}, timeout=120)
    if not resp:
        return None
    try:
        return resp['candidates'][0]['content']['parts'][0]['text']
    except Exception:
        print("❌ Gemini 返回结构异常")
        return None

def call_local(prompt, model_name=None, temperature=0.6):
    """
    直接调用 Gemini（正文写作），支持温度可调。
    temperature: 0.3-0.5(精准) / 0.6(均衡) / 0.7-0.9(创意)
    使用 gemini-3-pro-preview 获得最优的内容生成质量
    """
    # 附加图表类型参考信息
    chart_type_hint = """
    
    【图表类型速查表 - 根据内容自动选择最合适的图表】
    
    1. bar (柱状图) - 场景：对比分析、排名、市场份额对标
       示例：不同厂商的市场份额对比、2024-2025竞品对标
    
    2. line (折线图) - 场景：趋势分析、时序变化、增长率
       示例：2020-2024市场规模增长趋势、产能利用率变化
    
    3. pie (饼图) - 场景：占比分析、市场构成、成本分配
       示例：市场份额分布（A企业50%、B企业30%...）
    
    4. radar (雷达图) - 场景：多维评价、技术对标、能力评分
       示例：产品性能评分（可靠性80、效率90、成本70...）
    
    5. area (面积图) - 场景：堆积趋势、成本分解、层级分析
       示例：电力成本分布（燃煤60%、水电30%、风电10%）的变化趋势
    
    6. scatter (散点图) - 场景：相关性分析、分布显示、离散度
       示例：产品价格 vs 市场占有率的分布关系
    
    7. bubble (气泡图) - 场景：三维对比、企业规模、体量差异
       示例：不同企业的营收、利润、市场份额三维对比
    
    8. stacked_bar (堆积柱状图) - 场景：结构分析、层级分解
       示例：收入结构（主营业务、投资收益、其他）的年度对比
    
    9. heatmap (热力图) - 场景：矩阵分析、产业链、密度分布
       示例：产业链各环节的价值贡献度矩阵
    
    10. mixed (混合图) - 场景：多指标、复合分析、量纲不同
        示例：销量（柱状）+ 增长率（折线）的组合展示
    
    【选型原则】
    ✓ 每个子节最多用1个图表
    ✓ 同一份报告中避免重复使用相同图表类型
    ✓ 选择最能体现数据特征的图表类型
    ✓ 优先考虑 radar, area, heatmap 等复杂类型
    """
    
    model_id = model_name or CONF.GEMINI_GEN_MODEL
    # gemini-3-pro-preview 需要 v1beta API
    api_version = "v1beta"
    
    url = f"https://generativelanguage.googleapis.com/{api_version}/models/{model_id}:generateContent?key={CONF.GEMINI_API_KEY}"
    payload = {
        "contents": [{"parts": [{"text": prompt + chart_type_hint}]}],
        "generationConfig": {
            "temperature": temperature,
            "top_p": 0.95,  # 增加多样性同时保持相关性
            "top_k": 40
        }
    }
    resp = call_api_robust(url, payload, headers={"Content-Type": "application/json"}, proxies=CONF.PROXIES_CLOUD, timeout=120)
    if not resp:
        return None
    try:
        text = resp["candidates"][0]["content"]["parts"][0]["text"]
        return text.strip()
    except Exception as e:
        print(f"❌ 解析 Gemini 响应失败: {e}")
        return None

# ================== 🧭 最佳性价比流水线（Flash vs Pro 分工） ==================

def extract_first_json_block(text):
    """从任意文本中提取第一个 JSON 对象"""
    if not text:
        return None
    try:
        return json.loads(text)
    except Exception:
        pass
    match = re.search(r'(\{[\s\S]*\})', text)
    if match:
        try:
            return json.loads(match.group(1))
        except Exception:
            return None
    return None

def extract_code_block(text, language="python"):
    """提取指定语言的代码块"""
    if not text:
        return None
    pattern = rf"```{language}[\s\S]*?```"
    match = re.search(pattern, text, flags=re.IGNORECASE)
    if match:
        code = re.sub(rf"```{language}", "", match.group(0), flags=re.IGNORECASE)
        code = code.replace("```", "").strip()
        return code
    match = re.search(r"```[\s\S]*?```", text)
    if match:
        code = match.group(0).replace("```", "").strip()
        return code
    return None

def call_flash_json(prompt, temperature=0.25, task_type="heavy_reading"):
    """Flash JSON 输出（用于数据整理）"""
    return call_model(prompt, get_model(task_type), temperature=temperature, response_mime_type="application/json")

# ================== 🏛️ 全局上下文注入 ==================

def get_key_constraints():
    """固定的核心约束，供 Planner/Editor 参考"""
    return "\n".join([
        "- 严禁虚构数据，数字需可追溯",
        "- 每小节至少1表1图，图表数据需自洽",
        f"- 质量评分阈值 >= {CONF.QUALITY_THRESHOLD}",
        "- 语气专业、凝练，不要营销腔",
        "- Markdown 输出，不写代码/JSON（除图表数据生成阶段）"
    ])

def generate_global_thesis(topic, outline):
    """生成全书核心主旨（Global Thesis）"""
    prompt = f"""
你是总编辑，请凝练《{topic}》的全书核心主旨（80-120字）。
可用大纲: {json.dumps(outline, ensure_ascii=False)[:1200]}
要求: 用1段话写出全书核心论点和价值，不列点。
"""
    thesis = call_pro(prompt, temperature=0.35)
    return thesis or ""

def build_executive_summary(topic, context_mgr, last_chapter_title, global_thesis):
    """
    双结构摘要：
    - Rolling Window: 上一章 120-180 字
    - Global Thesis: 80-120 字
    """
    if not last_chapter_title or not context_mgr.generated_sections.get(last_chapter_title):
        return global_thesis or ""
    
    last_sections = context_mgr.generated_sections.get(last_chapter_title, {})
    last_text = "\n\n".join(list(last_sections.values()))
    
    prompt = f"""
你是 Executive Editor，生成下一章的上下文摘要，分为两段：
1) Rolling Window（上一章关键发现，120-180字）：强调数据/结论，为后续章节铺垫，避免重复上一章结尾表述。
2) Global Thesis（80-120字）：重申全书主旨，防止跑题。

主题: {topic}
上一章: {last_chapter_title}
上一章内容（节选）:
{last_text[:1200]}

全书主旨（参考）:
{global_thesis[:400]}

指令: 请结合全书核心目标，总结上一章的关键发现，并为下一章展开做铺垫。直接输出两段正文，不要标题。
"""
    summary = call_pro(prompt, temperature=0.35)
    return summary or global_thesis or ""

def write_checkpoint(checkpoint_path, chapter_index, chapter_title, executive_summary, global_thesis):
    """保存断点：章节进度 + 摘要"""
    payload = {
        "last_completed_chapter_index": chapter_index,
        "last_completed_chapter_title": chapter_title,
        "executive_summary": executive_summary,
        "global_thesis": global_thesis,
        "timestamp": datetime.now().isoformat()
    }
    try:
        with open(checkpoint_path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"⚠️ 断点保存失败: {e}")

def generate_style_guide(topic):
    """[Master Planner - Pro] 输出风格指南并缓存"""
    prompt = f"""
你是 [Master Planner - Pro]，为主题《{topic}》制定风格指南。
输出 Markdown 要点，覆盖：
- 语调: 专业、紧凑、避免营销腔
- 结构: 数据先行、逻辑递进、结论清晰
- 数据: 保留来源、禁止虚构、表格/图表各 1
- 语言: 避免重复，减少形容词，使用行业术语
- 图表: 只保留 1 张核心图，描述需简洁
请控制在 200-300 字，直接输出 Markdown 列表。
"""
    return call_pro(prompt, temperature=0.35)

def plan_section_flash(topic, chapter_title, section_title, subsections, style_guide, executive_summary="", key_constraints=""):
    """[Section Planner - Pro] 拆解章节任务（优先质量）"""
    subsections_text = ", ".join(subsections[:6])
    prompt = f"""
[Section Planner - Pro]
主题: {topic}
章节: {chapter_title} / {section_title}
小节: {subsections_text}
风格指南摘录: {style_guide[:600]}
Executive Summary: {executive_summary[:800]}
Key Constraints:
{key_constraints}

任务: 拆解本章节写作，给出：
- 数据要点（3-5 条，含指标/时间范围）
- 建议搜索词（4-6 个中英文混合）
- 建议图表（类型 + 维度 + 期望数据列）

要求: Markdown 列表，简短可执行，不要空话。
"""
    return call_pro(prompt, temperature=0.35)

def writer_flash_draft(topic, chapter_title, section_title, subsection_title, local_ctx, web_ctx, related_ctx, section_plan, style_guide):
    """[Writer - Flash] 粗写干巴巴初稿"""
    prompt = f"""
[Writer - Flash] 角色：蓝领采编，负责搜索/阅读/填充，输出干巴巴但准确的初稿。
主题: {topic}
位置: {chapter_title} / {section_title} / {subsection_title}
章节拆解: {section_plan}
风格指南: {style_guide[:800]}

[本地资料]
{local_ctx if local_ctx else "（无）"}

[网络资料]
{web_ctx if web_ctx else "（无）"}

[相关上下文]
{related_ctx if related_ctx else "（无）"}

写作要求:
1) 严禁虚构数据，仅引用资料中的数字/结论。
2) 结构: 关键数据要点 -> 分析解释 -> 风险/限制 -> 小结。
3) Markdown，使用 2-3 级标题，至少 1 个表格列出核心数字。
4) 引用规范：本地/网络/相关上下文中含有 [来源路径] 面包屑（filename > H1 > H2 > H3），引用关键数据时在正文点名来源；可选使用 Markdown 脚注 [^1]，脚注内容写“来源路径”。
5) 如果数据时间较旧（例如来源中年份早于需求），请在正文提示“数据时效性”。
6) 文风干练、直接，不做华丽措辞，保持“工程汇报”口吻。
7) 不要输出 JSON/代码/图表块，纯文字 + 表格即可。
"""
    return call_flash(prompt, temperature=0.55, task_type="heavy_reading")

def writer_flash_chart_data(topic, subsection_title, local_ctx, web_ctx, flash_draft, section_plan, style_guide):
    """[Writer - Flash] 提炼绘图数据"""
    prompt = f"""
[Writer - Flash] 只负责提炼绘图数据，不写代码。
主题: {topic}
子节: {subsection_title}
章节拆解: {section_plan}
风格指南: {style_guide[:400]}

[可用资料]
本地: {local_ctx if local_ctx else "（无）"}
网络: {web_ctx if web_ctx else "（无）"}
初稿: {flash_draft if flash_draft else "（初稿为空）"}

请给出单一 JSON 对象，字段: chart_type, title, x_label, y_label, data.labels[], data.datasets[].label/values, source。
规则: 
- 仅使用已出现的数据，不要猜测；若数据不足则返回空对象 {{}}。
- 保证 labels 和 values 数量一致，values 为数字。
- 限制 1 个图表。
"""
    raw = call_flash_json(prompt, temperature=0.25, task_type="heavy_reading")
    chart_data = extract_first_json_block(raw)
    if not chart_data:
        return None
    ok, errs = validate_json_chart_data(chart_data)
    if not ok:
        print(f"      ⚠️ Flash 图表数据校验失败: {errs}")
        return None
    return chart_data

def writer_pro_chart_data(topic, subsection_title, local_ctx, web_ctx, flash_draft, section_plan, style_guide):
    """[Writer - Pro] 兜底提炼绘图数据（质量优先）"""
    prompt = f"""
[Writer - Pro] 从现有资料中提炼一份图表 JSON。
主题: {topic}
子节: {subsection_title}
章节拆解: {section_plan}
风格指南: {style_guide[:400]}

[可用资料]
本地: {local_ctx if local_ctx else "（无）"}
网络: {web_ctx if web_ctx else "（无）"}
初稿: {flash_draft if flash_draft else "（初稿为空）"}

输出: 单一 JSON 对象，字段: chart_type, title, x_label, y_label, data.labels[], data.datasets[].label/values, source。
规则: 
- 仅使用已出现的数据，不要猜测；不足则返回空对象 {{}}。
- labels 与 values 数量一致，values 为数字。
- 限制 1 个图表。
"""
    raw = call_model(prompt, get_model("deep_thinking"), temperature=0.25, response_mime_type="application/json")
    chart_data = extract_first_json_block(raw)
    if not chart_data:
        return None
    ok, errs = validate_json_chart_data(chart_data)
    if not ok:
        print(f"      ⚠️ Pro 图表数据校验失败: {errs}")
        return None
    return chart_data

def execute_pro_chart_code(code_block, chart_data, output_path):
    """执行 Pro 生成的绘图代码，限制可用全局变量"""
    safe_globals = {
        "__builtins__": {
            "abs": abs, "min": min, "max": max, "range": range, "len": len,
            "float": float, "int": int, "sum": sum, "enumerate": enumerate, "round": round, "zip": zip,
            "list": list, "dict": dict, "str": str, "tuple": tuple
        },
        "plt": plt,
        "json": json,
        "math": __import__('math')  # 提供内置 math 库
    }
    try:
        import numpy as np
        safe_globals["np"] = np
    except ImportError:
        # numpy 不可用，代码中可能会失败，但至少不会导致程序崩溃
        print("      ℹ️ numpy 未安装，如果生成的代码需要 numpy 会失败")
    
    safe_globals["chart_data"] = chart_data
    safe_globals["output_path"] = output_path
    local_env = {}
    try:
        exec(code_block, safe_globals, local_env)
        render_fn = local_env.get("render") or safe_globals.get("render")
        if callable(render_fn):
            render_fn(chart_data, output_path)
        plt.close('all')
        return os.path.exists(output_path)
    except Exception as e:
        print(f"      ⚠️ Pro 绘图代码执行失败: {e}")
        return False

def writer_pro_chart(chart_data, output_dir, section_title, chapter_title):
    """[Writer - Pro] 根据整理好的数据写 Python 绘图代码并执行"""
    ok, errs = validate_json_chart_data(chart_data)
    if not ok:
        print(f"      ⚠️ 图表数据无效，跳过绘图: {errs}")
        return None, None
    
    safe_title = re.sub(r'[^\w\-]', '_', f"{chapter_title}_{section_title}")[:40]
    charts_dir = os.path.join(output_dir, "charts")
    os.makedirs(charts_dir, exist_ok=True)
    output_path = os.path.join(charts_dir, f"{safe_title}_{int(time.time())}.png")
    
    prompt = f"""
[Writer - Pro] 角色：只读已整理好的数据，写 Python 绘图代码。
输入 chart_data (JSON)：{json.dumps(chart_data, ensure_ascii=False)}
要求:
- 只用 matplotlib / numpy，不访问网络/文件系统，不调用系统命令。
- 定义 render(chart_data, output_path) 并在代码末尾调用它。
- 保存为 PNG 到 output_path，风格商务简洁，可读性高。
- 只输出一个 ```python``` 代码块，勿输出其他文字。
"""
    code_resp = call_pro(prompt, temperature=0.2)
    code_block = extract_code_block(code_resp, "python")
    
    success = False
    if code_block:
        success = execute_pro_chart_code(code_block, chart_data, output_path)
    
    if not success:
        print("      ↘️ 使用回退绘图方案")
        fallback_json = f"```json\n{json.dumps(chart_data, ensure_ascii=False)}\n```"
        success = create_chart_from_description_plotly(chart_data, output_path) or create_chart_from_description(fallback_json, output_path)
    
    return (output_path if success else None), code_block

def editor_pro_upgrade(topic, chapter_title, section_title, subsection_title, flash_draft, style_guide, section_plan, chart_data, chart_image_path):
    """[Editor - Pro] 升维润色，输出最终成品"""
    chart_ref = ""
    if chart_image_path:
        chart_ref = f"![{chart_data.get('title', '图表')}]({chart_image_path})"
    prompt = f"""
[Editor - Pro] 角色：白领审美+思考，对 Flash 初稿做“升维打击”。
主题: {topic}
位置: {chapter_title} / {section_title} / {subsection_title}
风格指南: {style_guide[:800]}
章节拆解: {section_plan}
图表: {json.dumps(chart_data, ensure_ascii=False) if chart_data else "（无图表）"} | 引用: {chart_ref if chart_ref else "无"}

[Flash 初稿]
{flash_draft if flash_draft else "（初稿为空）"}

任务:
1) 保留事实与数字，增强逻辑递进和行业洞察，修正语病。
2) 加入过渡句和结论，突出关键指标，适当补充背景。
3) 若有图表，正文中嵌入一次 Markdown 引用并给出一句解读：{chart_ref if chart_ref else "无"}。
4) 引用规范：正文引用关键数据时点名“来源路径”（来自 Flash 上下文的 [来源路径]: filename > H1 > H2 > H3），可使用 Markdown 脚注 [^1]，脚注内容写来源路径。
5) 若发现数据时效性不足（旧年份），需在文中提醒“数据时效性”。
6) 目标 900-1100 字，Markdown，避免再写 JSON/代码/大纲，不要添加额外 ##/### 标题（外层会包裹）。
7) 语气: 专业、凝练、可复用。
"""
    return call_pro(prompt, temperature=0.45)

# ================== 🚀 业务流程 ==================

def main():
    print("==========================================")
    print("   🏭 V23.0 工程重构版研报工厂            ")
    print("==========================================")
    
    # 1. 配置自检
    CONF.validate()
    
    # 2. 检查依赖库可用性（可选库，缺失时不中断流程）
    print("\n📦 检查可选依赖库...")
    
    # 检查 sklearn
    try:
        import sklearn
        print("✅ sklearn 已安装，将启用向量检索")
    except ImportError:
        print("ℹ️ sklearn 未安装，将使用关键词匹配进行检索")
        print("   💡 可选：安装 sklearn 获得更好的检索效果")
        print("   pip install scikit-learn")
    
    # 检查 numpy
    try:
        import numpy
        print("✅ numpy 已安装，支持所有图表类型")
    except ImportError:
        print("ℹ️ numpy 未安装，部分高级图表类型（如热力图）可能不可用")
        print("   💡 可选：安装 numpy 获得完整的图表功能")
        print("   pip install numpy")
    
    # 检查 Plotly 可用性
    try:
        import plotly
        print("✅ Plotly 已安装，将生成交互式图表")
    except ImportError:
        print("⚠️ Plotly 未安装，将使用 Matplotlib 生成静态图表")
        print("   💡 可选：安装 Plotly 获得更好的交互体验")
        print("   pip install plotly kaleido")
    
    # 检查是否需要自动调整结构
    if CONF.TARGET_PAGES > 0:
        print(f"\n📏 用户设定目标页数: {CONF.TARGET_PAGES} 页")
        chapters, sections, subsections = CONF.calculate_outline_structure()
        print(f"   └─ 自动优化结构: {chapters} 章 × {sections} 小节 × {subsections} 子节")
        print(f"   └─ 预计字数: {CONF.TARGET_PAGES * CONF.WORDS_PER_PAGE} 字")
        CONF.OUTLINE_CHAPTERS = chapters
        CONF.OUTLINE_SECTIONS = sections
        CONF.OUTLINE_SUBSECTIONS = subsections
    
    # 估算最终页数
    estimated_pages = CONF.estimate_page_count()
    total_subsections = CONF.OUTLINE_CHAPTERS * CONF.OUTLINE_SECTIONS * CONF.OUTLINE_SUBSECTIONS
    print(f"\n📊 当前生成参数配置:")
    print(f"   • 结构: {CONF.OUTLINE_CHAPTERS}章 x {CONF.OUTLINE_SECTIONS}节 x {CONF.OUTLINE_SUBSECTIONS}子节")
    print(f"   • 总任务: {total_subsections} 个写作单元")
    print(f"   • 预估产出: 约 {estimated_pages:.1f} 页 (按 {CONF.WORDS_PER_PAGE}字/页 计算)")
    print("☁️ 本轮使用 Gemini 3.0 Pro Preview（大纲 + 写作）- 最新统一架构")

    # 3. 获取主题
    topic = input("\n👉 请输入研报主题 (例如: 电力现货市场): ").strip()
    if not topic: return
    
    # 3. 路径准备 - 按主题自动创建独立文件夹
    folder_name = topic.replace(" ", "_").replace("/", "_")
    
    # 为该主题创建专属文件夹结构
    topic_base_dir = os.path.join(CONF.BASE_DIR, folder_name)
    input_dir = os.path.join(topic_base_dir, "materials")  # 素材存放目录
    output_dir = os.path.join(topic_base_dir, "output")    # 生成结果目录
    search_cache_dir = os.path.join(topic_base_dir, "search_cache")  # 搜索缓存目录
    
    # 自动创建文件夹结构
    for directory in [input_dir, output_dir, search_cache_dir]:
        if not os.path.exists(directory):
            os.makedirs(directory)
    
    print(f"📁 为主题 '{topic}' 创建独立工作目录:")
    print(f"   └─ 素材目录: {input_dir}")
    print(f"   └─ 输出目录: {output_dir}")
    print(f"   └─ 搜索缓存: {search_cache_dir}")
    
    # 4. 初始化知识库 & 上下文
    kb = MaterialManager(input_dir)
    context_mgr = ContextManager()
    
    # 5. Master Planner - Pro：大纲 & 风格指南，写入缓存
    print(f"\n💎 [Master Planner - Pro] 生成全书大纲 & 风格指南 -> Context Cache")
    outline_path = os.path.join(output_dir, "Structure.json")
    style_path = os.path.join(output_dir, "Style_Guide.md")
    cache_path = os.path.join(output_dir, "ContextCache.json")
    checkpoint_path = os.path.join(output_dir, "Checkpoint.json")
    
    if os.path.exists(outline_path):
        print("📋 加载现有大纲...")
        with open(outline_path, 'r') as f: outline = json.load(f)
    else:
        prompt = f"""
        Role: Chief Industry Expert in {topic}.
        Task: Design a rigorous, professional research report outline.
        Target Audience: High-level executives and technical directors.
        
        Requirements:
        1. Logical Flow: The structure must follow a logical progression (e.g., Market Analysis -> Technical Architecture -> Implementation Strategy -> Risk Control).
        2. Depth: Ensure deep vertical coverage of specific technologies/policies, not just surface-level breadth.
        3. Structure: {CONF.OUTLINE_CHAPTERS} Chapters -> {CONF.OUTLINE_SECTIONS} Sections -> {CONF.OUTLINE_SUBSECTIONS} Subsections.
        
        Output JSON: {{ "title": "...", "chapters": [ {{ "title": "...", "sections": [ {{ "title": "...", "subsections": ["..."] }} ] }} ] }}
        """
        res = call_gemini(prompt, json_mode=True)
        if not res:
            print("❌ 大纲生成失败，程序终止。")
            return
        try:
            clean_res = res.replace("```json", "").replace("```", "").strip()
            try:
                outline = json.loads(clean_res)
            except json.JSONDecodeError:
                match = re.search(r'(\{[\s\S]*\})', res)
                if match:
                    clean_json = match.group(1)
                    outline = json.loads(clean_json)
                else:
                    raise ValueError("无法提取 JSON 对象")
            with open(outline_path, "w") as f: json.dump(outline, f, indent=2, ensure_ascii=False)
        except Exception as e:
            print(f"❌ 大纲 JSON 解析失败: {e}")
            print(f"🔍 原始响应片段: {res[:500]}...")
            return
    if os.path.exists(style_path):
        with open(style_path, "r", encoding="utf-8") as f:
            style_guide = f.read()
        print("🎨 已加载风格指南")
    else:
        style_guide = generate_style_guide(topic) or ""
        with open(style_path, "w", encoding="utf-8") as f:
            f.write(style_guide)
        print("🎨 新生成风格指南")
    global_thesis = context_mgr.get_global_thesis()
    if not global_thesis:
        global_thesis = generate_global_thesis(topic, outline)
        context_mgr.set_global_thesis(global_thesis)
    context_mgr.set_master_plan(outline, style_guide)
    try:
        with open(cache_path, "w", encoding="utf-8") as f:
            json.dump({
                "outline": outline,
                "style_guide": style_guide,
                "global_thesis": global_thesis,
                "generated_at": datetime.now().isoformat()
            }, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

    # 6. Flash / Pro 写作流水线
    print(f"\n🏠 [Flash x Pro] 启动蓝领-白领流水线...")
    full_book = f"# {outline.get('title', topic)}\n\n"
    failed_sections = []
    any_content = False
    key_constraints = get_key_constraints()
    for i, chap in enumerate(outline.get('chapters', [])):
        chap_title = chap.get('title', f"Chapter {i+1}")
        full_book += f"# {chap_title}\n\n"
        print(f"\n📖 {chap_title}")
        for sec in chap.get('sections', []):
            sec_title = sec.get('title', 'Section')
            full_book += f"## {sec_title}\n\n"
            print(f"   📑 {sec_title}")
            sec_file = os.path.join(output_dir, f"{chap_title[:2]}_{sec_title[:5]}.md".replace(" ", "_").replace("/","-"))
            last_chapter_title = outline.get('chapters', [])[i-1].get('title') if i > 0 else None
            rolling_summary = build_executive_summary(topic, context_mgr, last_chapter_title, global_thesis)
            context_mgr.set_last_exec_summary(rolling_summary)
            section_plan = plan_section_flash(topic, chap_title, sec_title, sec.get('subsections', []), style_guide, rolling_summary, key_constraints)
            if section_plan:
                context_mgr.add_section_plan(f"{chap_title}/{sec_title}", section_plan)
            if os.path.exists(sec_file):
                print("      ✅ 已存在，跳过")
                with open(sec_file, 'r') as f: full_book += f.read() + "\n\n"
                continue
            sec_content = ""
            for sub in sec.get('subsections', []):
                print(f"      ✍️ [Flash] 撰写: {sub} ...")
                local_ctx = kb.retrieve(f"{topic} {sub}")
                related_ctx = context_mgr.get_related_context(f"{topic} {sub}")
                search_query = f"{topic} {sub} 数据 分析 现状"
                web_ctx = search_web(search_query, force=False, cache_dir=search_cache_dir)
                if not web_ctx and (not local_ctx or len(local_ctx) < 500):
                    web_ctx = search_web(f"{topic} {sub}", force=True, cache_dir=search_cache_dir)
                if not local_ctx and not web_ctx:
                    print(f"      ⚠️ 无外部资料命中（中英文都无），跳过此节点")
                    sec_content += f"### {sub}\n\n**⚠️ 数据缺失**\n\n本节点缺少相关的外部数据源（搜索无结果），为避免虚构内容，暂不生成。请手动补充资料或调整主题范围。\n\n"
                    failed_sections.append(f"{chap_title} / {sec_title} / {sub}")
                    continue
                flash_draft = writer_flash_draft(topic, chap_title, sec_title, sub, local_ctx, web_ctx, related_ctx, section_plan, style_guide)
                chart_data = writer_flash_chart_data(topic, sub, local_ctx, web_ctx, flash_draft, section_plan, style_guide)
                if not chart_data:
                    chart_data = writer_pro_chart_data(topic, sub, local_ctx, web_ctx, flash_draft, section_plan, style_guide)
                chart_path = None
                chart_rel_path = ""
                if chart_data:
                    chart_path, _ = writer_pro_chart(chart_data, output_dir, sub, chap_title)
                if chart_path:
                    chart_rel_path = os.path.relpath(chart_path, output_dir)
                    print(f"      🖼️ 图表生成完成: {chart_rel_path}")
                final_body = editor_pro_upgrade(topic, chap_title, sec_title, sub, flash_draft, style_guide, section_plan, chart_data, chart_rel_path)
                if not final_body and flash_draft:
                    final_body = flash_draft
                if not final_body:
                    print(f"      ❌ 本子节生成失败")
                    failed_sections.append(f"{chap_title} / {sec_title} / {sub}")
                    continue
                if chart_rel_path and chart_rel_path not in final_body:
                    final_body += f"\n\n![{chart_data.get('title', '图表')}]({chart_rel_path})\n"
                quality_score, feedback, _ = evaluate_content_quality(final_body, topic, sub)
                print(f"      📊 质量评分: {quality_score:.1f}/10 | {feedback[:60]}")
                if quality_score < CONF.QUALITY_THRESHOLD:
                    print("      🔁 质量未达标，切换 Pro 再润色一轮")
                    retry_body = editor_pro_upgrade(topic, chap_title, sec_title, sub, final_body, style_guide, section_plan, chart_data, chart_rel_path)
                    if retry_body:
                        final_body = retry_body
                        quality_score, feedback, _ = evaluate_content_quality(final_body, topic, sub)
                        print(f"      📊 二次质量评分: {quality_score:.1f}/10 | {feedback[:60]}")
                sec_content += f"### {sub}\n\n{final_body}\n\n"
                any_content = True
                context_mgr.add_section(chap_title, sub, final_body)
            if sec_content:
                with open(sec_file, "w", encoding="utf-8") as f: f.write(sec_content)
                full_book += sec_content
            else:
                print("      ❌ 本节未生成有效内容")
        # 章节完成后保存断点（便于续跑）
        write_checkpoint(checkpoint_path, i, chap_title, context_mgr.get_last_exec_summary(), global_thesis)
    # 8. 装订
    if not any_content:
        print("\n❌ 未生成任何有效章节，已停止装订。请检查本地模型或资料。")
        if failed_sections:
            print("未完成列表:")
            for item in failed_sections: print(f" - {item}")
        return

    final_path = os.path.join(output_dir, "Final_Book.md")
    with open(final_path, "w", encoding="utf-8") as f: f.write(full_book)
    print(f"\n🎉🎉🎉 任务完成！文件: {final_path}")
    
    # 打印生成统计
    print(f"\n📊 最终生成统计:")
    print(f"   已生成章节数: {len(context_mgr.generated_sections)}")
    
    total_sections_gen = sum(len(secs) for secs in context_mgr.generated_sections.values())
    
    # 计算实际字数
    total_chars = 0
    for ch_content in context_mgr.generated_sections.values():
        for sec_content in ch_content.values():
            total_chars += len(sec_content)
            
    actual_pages = total_chars / CONF.WORDS_PER_PAGE
    
    print(f"   已生成小节数: {total_sections_gen}")
    print(f"   总字数: {total_chars:,} 字")
    print(f"   实际页数: {actual_pages:.1f} 页 (目标: {CONF.TARGET_PAGES} 页)")
    print(f"   失败小节数: {len(failed_sections)}")
    
    print(f"\n{context_mgr.get_summary()}")
    
    if failed_sections:
        print("⚠️ 以下子节写作失败，请手动补写或重跑:")
        for item in failed_sections: print(f" - {item}")

    # 自动转换 Word（带样式）
    docx_name = f"{folder_name}_Research_Report.docx"
    docx_path = os.path.join(output_dir, docx_name)
    ref_doc = get_reference_doc_path()
    if ref_doc and not os.path.exists(ref_doc):
        print(f"⚠️ 未找到样式母版 {ref_doc}，将使用 Pandoc 默认样式。")
        ref_doc = None
    convert_md_to_docx(
        md_filename=final_path,
        output_filename=docx_path,
        reference_doc=ref_doc,
        resource_path=output_dir
    )
    print(f"📄 已生成 Word 文档: {docx_path}")

# ============ Word 转换工具 ============ #

def get_reference_doc_path():
    """返回样式母版路径，默认放在 BASE_DIR 下"""
    return os.path.join(CONF.BASE_DIR, "reference.docx")

def convert_md_to_docx(md_filename, output_filename, reference_doc=None, resource_path="."):
    """
    将 Markdown 转换为排版完美的 Word，支持样式母版与资源路径。
    """
    print(f"🔄 正在将 {md_filename} 转换为 Word 文档...")
    extra_args = [
        '--toc',
        '--toc-depth=3',
        f'--resource-path={resource_path}',
        '--standalone'
    ]
    if reference_doc and os.path.exists(reference_doc):
        extra_args.append(f'--reference-doc={reference_doc}')
    try:
        import pypandoc
        pypandoc.convert_file(
            md_filename,
            'docx',
            outputfile=output_filename,
            extra_args=extra_args
        )
        print(f"✅ 转换成功！文档已生成: {output_filename}")
    except Exception as e:
        print(f"❌ 转换失败: {e}")
        print(f"建议手动执行命令: pandoc {md_filename} -o {output_filename} {'--reference-doc='+reference_doc if reference_doc else ''} --toc")

if __name__ == "__main__":
    main()
